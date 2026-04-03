"""
Professional Export Services for DidactAI

Generates clean, professional exam and quiz documents without question type labels.
This module provides services for exporting educational content to various formats
including PDF, DOCX, HTML, and ZIP archives.
"""

import io
import os
import json
import zipfile
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path

from django.conf import settings
from django.template.loader import render_to_string
from django.utils import timezone
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage

# Optional image processing for efficient logo embedding
try:
    from PIL import Image as PILImage
except ImportError:  # pragma: no cover
    PILImage = None

try:
    import cairosvg
except ImportError:  # pragma: no cover
    cairosvg = None

# Import export libraries
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Image
    from reportlab.lib.utils import ImageReader
    import reportlab.rl_config
    reportlab.rl_config.warnOnMissingFontGlyphs = 0  # Disable warnings for missing glyphs
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# WeasyPrint removed due to Windows GTK dependencies
# Focusing on ReportLab PDF and DOCX exports which work perfectly
WEASYPRINT_AVAILABLE = False

import logging

logger = logging.getLogger(__name__)


def _guess_logo_mime(name: str) -> str:
    ext = (os.path.splitext(name or "")[1] or "").lower()
    if ext == '.svg':
        return 'image/svg+xml'
    if ext in ['.jpg', '.jpeg']:
        return 'image/jpeg'
    if ext == '.png':
        return 'image/png'
    if ext == '.gif':
        return 'image/gif'
    return 'image/png'


def _read_logo_bytes_from_branding(branding: Optional[Dict[str, Any]]) -> Optional[Tuple[bytes, str]]:
    """Resolve the uploaded logo into bytes suitable for embedding.

    Prefers local filesystem paths when available, but can also read from Django storage
    (important when ImageField storage isn't local).

    Returns:
        (bytes, mime_type) or None
    """
    if not branding:
        return None

    candidates: List[Tuple[str, str]] = []

    logo_path = branding.get('logo_path')
    if logo_path:
        candidates.append(('path', str(logo_path)))

    logo_url = branding.get('logo_url')
    if logo_url and isinstance(logo_url, str) and logo_url.startswith('/media/'):
        rel = logo_url.replace('/media/', '').lstrip('/').replace('/', os.sep)
        candidates.append(('path', os.path.join(settings.MEDIA_ROOT, rel)))

    logo_filename = branding.get('logo_filename')
    if logo_filename:
        # May be a storage path like "export_logos/foo.png"
        candidates.append(('storage', str(logo_filename)))
        # Also try as a relative path under MEDIA_ROOT
        candidates.append(('path', os.path.join(settings.MEDIA_ROOT, str(logo_filename).replace('/', os.sep))))
        # And as a basename under export_logos/
        candidates.append(('path', os.path.join(settings.MEDIA_ROOT, 'export_logos', os.path.basename(str(logo_filename)))))

    for kind, val in candidates:
        try:
            if kind == 'path' and val and os.path.exists(val):
                with open(val, 'rb') as f:
                    return f.read(), _guess_logo_mime(val)
            if kind == 'storage' and val and default_storage.exists(val):
                with default_storage.open(val, 'rb') as f:
                    return f.read(), _guess_logo_mime(val)
        except Exception as e:
            logger.warning(f"Failed reading logo from {kind}={val}: {e}")
            continue

    return None


def _prepare_logo_for_embedding(logo_bytes: bytes, mime_type: str, max_px: int = 256) -> Tuple[bytes, str]:
    """Normalize the logo to a small embedded PNG (fast + consistent across exporters).

    - SVG: convert to PNG if cairosvg is available.
    - Raster: downscale to max_px and convert to PNG if Pillow is available.
    """
    if not logo_bytes:
        return logo_bytes, mime_type

    # SVG -> PNG
    if mime_type == 'image/svg+xml':
        if cairosvg is None:
            logger.warning("SVG logo provided but cairosvg is not installed; skipping SVG conversion.")
            return logo_bytes, mime_type
        try:
            png_bytes = cairosvg.svg2png(bytestring=logo_bytes, output_width=max_px, output_height=max_px)
            return png_bytes, 'image/png'
        except Exception as e:
            logger.warning(f"Failed converting SVG logo to PNG: {e}")
            return logo_bytes, mime_type

    # Raster resize/convert
    if PILImage is None:
        return logo_bytes, mime_type

    try:
        img = PILImage.open(io.BytesIO(logo_bytes))
        # Ensure deterministic output + broad compatibility
        img = img.convert('RGBA') if img.mode not in ('RGB', 'RGBA') else img
        img.thumbnail((max_px, max_px))
        out = io.BytesIO()
        img.save(out, format='PNG', optimize=True)
        return out.getvalue(), 'image/png'
    except Exception as e:
        logger.warning(f"Failed resizing/converting logo for embedding: {e}")
        return logo_bytes, mime_type


class PDFExporter:
    """Service for exporting content to PDF format using ReportLab"""
    
    def __init__(self):
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab is required for PDF export")
        self._setup_unicode_fonts()
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _build_logo_flowable(self, branding: Optional[Dict[str, Any]], max_w=None, max_h=None):
        """Return a ReportLab flowable for the uploaded logo (or None).

        Uses bytes-based loading so it works with non-local Django storage too.
        """
        try:
            if max_w is None:
                max_w = 1.2 * inch
            if max_h is None:
                max_h = 1.2 * inch
            resolved = _read_logo_bytes_from_branding(branding)
            if not resolved:
                return None
            logo_bytes, mime = resolved
            logo_bytes, _ = _prepare_logo_for_embedding(logo_bytes, mime, max_px=256)

            # Compute aspect-preserving size
            reader = ImageReader(io.BytesIO(logo_bytes))
            iw, ih = reader.getSize()
            if not iw or not ih:
                return None

            scale = min(max_w / float(iw), max_h / float(ih))
            w = float(iw) * scale
            h = float(ih) * scale
            return Image(io.BytesIO(logo_bytes), width=w, height=h)
        except Exception as e:
            logger.warning(f"Could not build logo flowable: {e}")
            return None
    
    def _setup_unicode_fonts(self):
        """Setup Unicode fonts that support Turkish characters"""
        try:
            # Try to register DejaVu fonts (widely available and support Turkish)
            from reportlab.lib.fonts import addMapping
            
            # Use Helvetica as default (has better Unicode support than Times)
            # ReportLab's built-in fonts should handle basic Turkish characters
            self.unicode_font_normal = 'Helvetica'
            self.unicode_font_bold = 'Helvetica-Bold'
            
            # Test if we can find DejaVu fonts (better Unicode support)
            import os
            possible_font_paths = [
                r'C:\Windows\Fonts\DejaVuSans.ttf',
                r'C:\Windows\Fonts\arial.ttf',
                r'C:\Windows\Fonts\calibri.ttf',
                '/System/Library/Fonts/Helvetica.ttc',
                '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
            ]
            
            font_found = False
            for font_path in possible_font_paths:
                if os.path.exists(font_path):
                    try:
                        if 'DejaVu' in font_path:
                            pdfmetrics.registerFont(TTFont('Unicode', font_path))
                            self.unicode_font_normal = 'Unicode'
                            font_found = True
                            logger.info(f"Registered Unicode font: {font_path}")
                            break
                        elif 'arial' in font_path.lower():
                            pdfmetrics.registerFont(TTFont('Arial', font_path))
                            self.unicode_font_normal = 'Arial'
                            font_found = True
                            logger.info(f"Registered Arial font: {font_path}")
                            break
                    except Exception as e:
                        logger.warning(f"Could not register font {font_path}: {e}")
                        continue
            
            if not font_found:
                logger.info("Using built-in Helvetica font (supports basic Turkish characters)")
                
        except Exception as e:
            logger.warning(f"Font setup error: {e}. Using default fonts.")
            self.unicode_font_normal = 'Helvetica'
            self.unicode_font_bold = 'Helvetica-Bold'
    
    def _setup_custom_styles(self):
        """Setup custom paragraph styles for professional academic documents"""
        # Title style - professional university appearance with Unicode support
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Title'],
            fontSize=22,
            spaceAfter=20,
            spaceBefore=0,
            alignment=TA_CENTER,
            textColor=colors.black,
            fontName=getattr(self, 'unicode_font_bold', 'Helvetica-Bold')
        ))
        
        # Subtitle for exam details
        self.styles.add(ParagraphStyle(
            name='Subtitle',
            parent=self.styles['Normal'],
            fontSize=12,
            spaceAfter=15,
            spaceBefore=5,
            alignment=TA_CENTER,
            textColor=colors.black,
            fontName=getattr(self, 'unicode_font_normal', 'Helvetica')
        ))
        
        # Header style - clean and professional
        self.styles.add(ParagraphStyle(
            name='CustomHeader',
            parent=self.styles['Heading1'],
            fontSize=14,
            spaceBefore=25,
            spaceAfter=15,
            textColor=colors.black,
            fontName=getattr(self, 'unicode_font_bold', 'Helvetica-Bold'),
            borderWidth=1,
            borderColor=colors.black,
            borderPadding=8,
            alignment=TA_CENTER
        ))
        
        # Instructions style with box
        self.styles.add(ParagraphStyle(
            name='Instructions',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceBefore=10,
            spaceAfter=20,
            leftIndent=15,
            rightIndent=15,
            fontName=getattr(self, 'unicode_font_normal', 'Helvetica'),
            borderWidth=1,
            borderColor=colors.grey,
            borderPadding=12,
            backColor=colors.lightgrey
        ))
        
        # Question style - justified and professional with compact spacing
        self.styles.add(ParagraphStyle(
            name='Question',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceBefore=15,  # Adequate space before question
            spaceAfter=8,    # Compact space after
            leftIndent=0,
            fontName=getattr(self, 'unicode_font_normal', 'Helvetica'),
            alignment=TA_JUSTIFY,
            leading=14,  # Line height
            keepWithNext=1  # Keep question with its content
        ))
        
        # Question number style
        self.styles.add(ParagraphStyle(
            name='QuestionNumber',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceBefore=18,  # Adequate space before question number
            spaceAfter=6,    # Space after question number
            fontName=getattr(self, 'unicode_font_bold', 'Helvetica-Bold'),
            keepWithNext=1  # Keep question number with question text
        ))
        
        # Option style for multiple choice - properly indented with compact spacing
        self.styles.add(ParagraphStyle(
            name='Option',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceBefore=3,  # Minimal space between options
            spaceAfter=3,   # Minimal space between options
            leftIndent=25,  # Indent options
            fontName=getattr(self, 'unicode_font_normal', 'Helvetica'),
            alignment=TA_JUSTIFY,
            leading=13,  # Compact line height
            keepTogether=1  # Keep all options together if possible
        ))
        
        # Answer lines style
        self.styles.add(ParagraphStyle(
            name='AnswerLines',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceBefore=8,
            spaceAfter=4,
            leftIndent=0,
            fontName=getattr(self, 'unicode_font_normal', 'Helvetica')
        ))
        
        # Answer style for answer keys
        self.styles.add(ParagraphStyle(
            name='Answer',
            parent=self.styles['Normal'],
            fontSize=10,
            leftIndent=20,
            textColor=colors.darkgreen,
            fontName=getattr(self, 'unicode_font_normal', 'Helvetica')
        ))
    
    def export_quiz(self, quiz_data: Dict[str, Any], branding: Dict[str, Any] = None) -> io.BytesIO:
        """Export quiz to PDF format with RDUU university design"""
        buffer = io.BytesIO()
        
        # Store branding data for header/footer use
        self.branding_data = branding or {}
        
        # Create PDF document with RDUU margins
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        story = []
        
        # Add RDUU cover page
        story.extend(self._create_rduu_cover_page(quiz_data, branding))
        
        # Add page break after cover page
        story.append(PageBreak())
        
        # Questions Section with RDUU formatting - start immediately after page break
        questions = quiz_data.get('questions', [])
        for i, question in enumerate(questions, 1):
            question_type = question.get('type', 'multiple_choice')
            points = question.get('points', 1)
            
            # Question header with RDUU style - improved page break control
            q_header = f"<b>Question {i}. ({points} point{'s' if points != 1 else ''})</b>"
            q_header_style = ParagraphStyle(
                name='RDUUQuestionHeader',
                parent=self.styles['Normal'],
                fontSize=11,
                spaceBefore=20,  # More space before question
                spaceAfter=8,
                fontName='Helvetica-Bold',
                keepWithNext=1  # Keep header with question text
            )
            story.append(Paragraph(q_header, q_header_style))
            
            # Question text with better formatting
            question_text = question.get('question', '')
            q_text_style = ParagraphStyle(
                name='RDUUQuestionText',
                parent=self.styles['Normal'],
                fontSize=11,
                spaceBefore=6,
                spaceAfter=10,
                fontName='Helvetica',
                alignment=TA_JUSTIFY,
                leading=14,  # Better line spacing
                keepWithNext=1  # Keep question text with its options
            )
            story.append(Paragraph(question_text, q_text_style))
            
            # Handle different question types with RDUU formatting
            if question_type == 'multiple_choice' and question.get('options'):
                # Ensure we have unique options and proper formatting
                unique_options = list(dict.fromkeys(question['options']))  # Remove duplicates
                # Support up to 5 options (A-E)
                option_style = ParagraphStyle(
                    name='RDUUOption',
                    parent=self.styles['Normal'],
                    fontSize=11,
                    spaceBefore=4,  # Small space between options
                    spaceAfter=4,   # Small space between options
                    leftIndent=30,  # Better indentation
                    fontName='Helvetica',
                    alignment=TA_JUSTIFY,
                    leading=13,  # Compact line spacing
                    keepTogether=1  # Prevent option from splitting across pages
                )
                
                # Add all options with proper page break control
                for j, option in enumerate(unique_options[:5]):  # Limit to 5 options max (A-E)
                    option_letter = chr(65 + j)  # A, B, C, D, E
                    option_text = f"{option_letter}. {option}"
                    option_para = Paragraph(option_text, option_style)
                    story.append(option_para)
            
            elif question_type == 'true_false':
                option_style = ParagraphStyle(
                    name='RDUUOption',
                    parent=self.styles['Normal'],
                    fontSize=11,
                    spaceBefore=3,
                    spaceAfter=3,
                    leftIndent=25,
                    fontName='Helvetica'
                )
                story.append(Paragraph("A. True", option_style))
                story.append(Paragraph("B. False", option_style))
            
            elif question_type == 'short_answer':
                story.append(Spacer(1, 12))
                answer_style = ParagraphStyle(
                    name='RDUUAnswer',
                    parent=self.styles['Normal'],
                    fontSize=11,
                    spaceBefore=8,
                    spaceAfter=4,
                    fontName='Helvetica-Bold'
                )
                story.append(Paragraph("Answer:", answer_style))
                story.append(Spacer(1, 8))
                line_style = ParagraphStyle(
                    name='RDUUAnswerLine',
                    parent=self.styles['Normal'],
                    fontSize=11,
                    spaceBefore=4,
                    spaceAfter=4,
                    fontName='Helvetica'
                )
                for _ in range(5):  # 5 lines for short answer
                    story.append(Paragraph("_" * 85, line_style))
            
            elif question_type == 'fill_blank':
                story.append(Spacer(1, 12))
                answer_style = ParagraphStyle(
                    name='RDUUFillAnswer',
                    parent=self.styles['Normal'],
                    fontSize=11,
                    spaceBefore=8,
                    spaceAfter=8,
                    fontName='Helvetica-Bold'
                )
                story.append(Paragraph("Answer: " + "_" * 60, answer_style))
            
            elif question_type == 'essay':
                story.append(Spacer(1, 12))
                essay_style = ParagraphStyle(
                    name='RDUUEssayAnswer',
                    parent=self.styles['Normal'],
                    fontSize=11,
                    spaceBefore=8,
                    spaceAfter=10,
                    fontName='Helvetica-Bold'
                )
                story.append(Paragraph("Answer: (Use the space below for your complete response)", essay_style))
                line_style = ParagraphStyle(
                    name='RDUUEssayLine',
                    parent=self.styles['Normal'],
                    fontSize=11,
                    spaceBefore=4,
                    spaceAfter=4,
                    fontName='Helvetica'
                )
                for _ in range(12):  # More lines for essay questions
                    story.append(Paragraph("_" * 85, line_style))
            
            # Add appropriate space between questions to prevent crowding
            story.append(Spacer(1, 16))  # Balanced space between questions
        
        # Build PDF with RDUU template
        doc.build(story, onFirstPage=self._add_rduu_header_footer, onLaterPages=self._add_rduu_header_footer)
        buffer.seek(0)
        return buffer
    
    def _add_header_footer(self, canvas, doc):
        """Add headers and footers to each page with enhanced information"""
        canvas.saveState()
        
        # Header line
        canvas.setStrokeColor(colors.black)
        canvas.setLineWidth(0.5)
        canvas.line(72, letter[1] - 50, letter[0] - 72, letter[1] - 50)
        
        # Enhanced footer with academic information
        canvas.setFont('Times-Roman', 9)
        
        # Left side: Academic year and semester
        if hasattr(self, 'branding_data') and self.branding_data:
            footer_left = []
            if self.branding_data.get('academic_year'):
                footer_left.append(f"Academic Year: {self.branding_data['academic_year']}")
            if self.branding_data.get('semester'):
                footer_left.append(self.branding_data['semester'])
            
            if footer_left:
                left_text = ' | '.join(footer_left)
                canvas.drawString(72, 30, left_text)
        
        # Center: Page number
        canvas.setFont('Times-Roman', 10)
        page_text = f"Page {doc.page}"
        text_width = canvas.stringWidth(page_text, 'Times-Roman', 10)
        canvas.drawString((letter[0] - text_width) / 2, 30, page_text)
        
        # Right side: University name (abbreviated)
        if hasattr(self, 'branding_data') and self.branding_data:
            if self.branding_data.get('university_name'):
                uni_name = self.branding_data['university_name']
                # Abbreviate long university names
                if len(uni_name) > 20:
                    words = uni_name.split()
                    uni_name = ''.join([word[0] for word in words if word]) + 'U'
                
                canvas.setFont('Times-Roman', 9)
                right_text_width = canvas.stringWidth(uni_name, 'Times-Roman', 9)
                canvas.drawString(letter[0] - 72 - right_text_width, 30, uni_name)
        
        # Footer line
        canvas.line(72, 50, letter[0] - 72, 50)
        
        # Add watermark if specified
        if hasattr(self, 'branding_data') and self.branding_data and self.branding_data.get('watermark'):
            watermark_text = self.branding_data['watermark']
            if watermark_text.strip():
                # Save current state
                canvas.saveState()
                
                # Set watermark properties
                canvas.setFont('Helvetica-Bold', 48)
                canvas.setFillColor(colors.lightgrey)
                canvas.setFillAlpha(0.3)  # Make it semi-transparent
                
                # Calculate center position
                page_width = letter[0]
                page_height = letter[1]
                
                # Rotate and draw watermark diagonally across the page
                canvas.translate(page_width/2, page_height/2)
                canvas.rotate(45)  # 45-degree angle
                
                # Calculate text width to center it
                text_width = canvas.stringWidth(watermark_text, 'Helvetica-Bold', 48)
                canvas.drawCentredString(0, 0, watermark_text)
                
                # Restore state
                canvas.restoreState()
        
        canvas.restoreState()
    
    def _create_rduu_cover_page(self, quiz_data: Dict[str, Any], branding: Dict[str, Any] = None) -> List:
        """Create professional cover page with proper text placement"""
        elements = []
        branding = branding or {}
        
        # Top header with academic info (removed RDUU)
        header_data = [[
            f"Academic Year: {branding.get('academic_year', '2025-2026')} | {branding.get('semester', 'Fall Semester')}",
            "Page 1"
        ]]
        header_table = Table(header_data, colWidths=[5*inch, 1.5*inch])
        header_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('ALIGN', (0, 0), (0, 0), 'LEFT'),
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        elements.append(header_table)
        elements.append(Spacer(1, 25))
        
        # University logo (embedded) - only shown when uploaded
        logo_flowable = self._build_logo_flowable(branding, max_w=1.4*inch, max_h=1.4*inch)
        if logo_flowable:
            logo_table = Table([[logo_flowable]], colWidths=[6*inch])
            logo_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            elements.append(logo_table)
            elements.append(Spacer(1, 12))
        else:
            elements.append(Spacer(1, 15))
        
        # University name with better positioning
        university_style = ParagraphStyle(
            name='UniversityHeader',
            parent=self.styles['Normal'],
            fontSize=20,
            spaceAfter=6,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold',
            textColor=colors.black,
            leading=22
        )
        university_name = branding.get('university_name', 'RAUF DENKTAS UNIVERSITY')
        elements.append(Paragraph(university_name, university_style))
        
        # Department info with tighter spacing
        dept_style = ParagraphStyle(
            name='DepartmentInfo',
            parent=self.styles['Normal'],
            fontSize=12,
            spaceAfter=10,
            alignment=TA_CENTER,
            fontName='Helvetica',
            leading=14
        )
        dept_info = f"{branding.get('department', 'SOFTWARE ENGINEERING')} | {branding.get('course', 'Software Engineering')}"
        elements.append(Paragraph(dept_info, dept_style))
        
        # Subject title with optimized spacing
        subject_style = ParagraphStyle(
            name='SubjectTitle',
            parent=self.styles['Normal'],
            fontSize=16,
            spaceAfter=15,
            spaceBefore=10,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold',
            leading=18
        )
        elements.append(Paragraph(quiz_data.get('subject', 'Quiz'), subject_style))
        
        # Instructor and date info box with tighter spacing
        instructor_info = f"Instructor: {branding.get('instructor', 'RAMS')} | Date: {branding.get('exam_date', '2025-11-07')}"
        info_box = Table([[instructor_info]], colWidths=[6*inch])
        info_box.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('BOX', (0, 0), (-1, -1), 1, colors.black),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        elements.append(info_box)
        elements.append(Spacer(1, 15))
        
        # Main exam title with compact spacing
        exam_title_style = ParagraphStyle(
            name='ExamTitle',
            parent=self.styles['Normal'],
            fontSize=18,
            spaceAfter=12,
            spaceBefore=5,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold',
            leading=20
        )
        exam_title = quiz_data.get('title', 'Examination')
        elements.append(Paragraph(exam_title, exam_title_style))
        
        # Duration and points with reduced spacing
        duration_info = f"Duration: {quiz_data.get('estimated_duration', '75 minutes')} | Total Points: {quiz_data.get('total_points', '50')}"
        duration_style = ParagraphStyle(
            name='DurationInfo',
            parent=self.styles['Normal'],
            fontSize=12,
            spaceAfter=15,
            alignment=TA_CENTER,
            fontName='Helvetica',
            leading=14
        )
        elements.append(Paragraph(duration_info, duration_style))
        
        # Instructions box with line spacing = 1 and grey background
        instruction_style = ParagraphStyle(
            name='InstructionText',
            parent=self.styles['Normal'],
            fontSize=11,
            leading=11,  # Line spacing = 1 (leading = fontSize)
            spaceAfter=2,
            fontName='Helvetica'
        )
        
        instruction_title_style = ParagraphStyle(
            name='InstructionTitle',
            parent=self.styles['Normal'],
            fontSize=12,
            leading=12,
            spaceAfter=6,
            fontName='Helvetica-Bold'
        )
        
        instructions_content = [
            [Paragraph("INSTRUCTIONS:", instruction_title_style)],
            [Paragraph("• Read each question carefully and completely before answering", instruction_style)],
            [Paragraph("• For multiple choice questions, select the best answer", instruction_style)],
            [Paragraph("• Write clearly and legibly for all written responses", instruction_style)],
            [Paragraph("• Show all work for calculation problems where applicable", instruction_style)],
            [Paragraph("• Review your answers before submitting", instruction_style)],
            [Paragraph("• Ask the instructor if you have any questions", instruction_style)]
        ]
        
        instructions_table = Table(instructions_content, colWidths=[6*inch])
        instructions_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
            ('BOX', (0, 0), (-1, -1), 1.5, colors.black),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 12),
            ('RIGHTPADDING', (0, 0), (-1, -1), 12),
        ]))
        elements.append(instructions_table)
        elements.append(Spacer(1, 20))
        
        # Student information section with compact layout
        student_fields = [
            ['Student Full Name:', '_' * 60],
            ['Student ID/Number:', '_' * 35],
            ['Date:', '_' * 20],
            ['Signature:', '_' * 40]
        ]
        
        student_table = Table(student_fields, colWidths=[2*inch, 4*inch])
        student_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),
            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        elements.append(student_table)
        
        return elements
    
    def _add_rduu_header_footer(self, canvas, doc):
        """Add professional headers and footers to each page"""
        canvas.saveState()
        
        # Only add header/footer to pages after the cover page
        if doc.page > 1:
            # Footer with academic info and page number (as requested)
            canvas.setFont('Helvetica', 10)
            footer_text = f"Academic Year: {self.branding_data.get('academic_year', '2025-2026')} | {self.branding_data.get('semester', 'Fall Semester')}    Page {doc.page}"
            
            # Draw footer at bottom of page
            canvas.drawString(72, 30, footer_text)
            
            # Top line for header separation
            canvas.setStrokeColor(colors.black)
            canvas.setLineWidth(0.5)
            canvas.line(72, letter[1] - 65, letter[0] - 72, letter[1] - 65)
            
            # Footer line
            canvas.line(72, 50, letter[0] - 72, 50)
        
        canvas.restoreState()
    
    def _create_single_cover_page(self, quiz_data: Dict[str, Any], branding: Dict[str, Any] = None) -> List:
        """Create a single comprehensive cover page with all necessary information"""
        elements = []
        
        # Enhanced University Header with better spacing
        if branding:
            # Add top spacing
            elements.append(Spacer(1, 30))
            
            # Add logo first if available
            if branding.get('logo_path') or branding.get('has_logo') or branding.get('logo_url'):
                try:
                    logo_path = None
                    
                    # Method 1: Direct path from branding settings
                    if branding.get('logo_path'):
                        potential_path = branding['logo_path']
                        if os.path.exists(potential_path):
                            logo_path = potential_path
                            logger.info(f"Found logo for cover page using direct path: {logo_path}")
                    
                    # Method 2: Convert URL to file path if direct path didn't work
                    if not logo_path and branding.get('logo_url'):
                        from django.conf import settings
                        logo_url = branding['logo_url']
                        if logo_url.startswith('/media/'):
                            relative_path = logo_url.replace('/media/', '').lstrip('/')
                            potential_path = os.path.join(settings.MEDIA_ROOT, relative_path)
                            if os.path.exists(potential_path):
                                logo_path = potential_path
                                logger.info(f"Found logo for cover page using URL conversion: {logo_path}")
                    
                    # Method 3: Try using filename from branding settings
                    if not logo_path and branding.get('logo_filename'):
                        from django.conf import settings
                        filename = branding['logo_filename']
                        potential_path = os.path.join(settings.MEDIA_ROOT, 'export_logos', os.path.basename(filename))
                        if os.path.exists(potential_path):
                            logo_path = potential_path
                            logger.info(f"Found logo for cover page using filename: {logo_path}")
                    
                    if logo_path and os.path.exists(logo_path):
                        from reportlab.platypus import Image as ReportLabImage
                        try:
                            # Create logo image with appropriate size
                            logo_img = ReportLabImage(logo_path, width=100, height=100)
                            
                            # Create a table to center the logo properly
                            logo_table = Table([[logo_img]], colWidths=[6*inch])
                            logo_table.setStyle(TableStyle([
                                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                            ]))
                            
                            elements.append(logo_table)
                            elements.append(Spacer(1, 15))
                            logger.info(f"Successfully added logo to cover page from: {logo_path}")
                        except Exception as img_error:
                            logger.error(f"Error creating logo image for cover page: {img_error}")
                except Exception as e:
                    logger.error(f"Could not add logo to cover page: {e}")
            
            # University name and logo area with improved typography
            if branding.get('university_name') or branding.get('institution_name'):
                institution = branding.get('university_name') or branding.get('institution_name')
                
                # University header style - larger and more prominent
                uni_style = ParagraphStyle(
                    name='UniversityName',
                    parent=self.styles['Normal'],
                    fontSize=24,  # Increased from 20
                    alignment=TA_CENTER,
                    spaceBefore=20,
                    spaceAfter=12,
                    textColor=colors.black,
                    fontName=getattr(self, 'unicode_font_bold', 'Helvetica-Bold'),
                    borderWidth=2,
                    borderColor=colors.black,
                    borderPadding=15
                )
                elements.append(Paragraph(institution.upper(), uni_style))
                
                # Department and course info
                dept_info = []
                if branding.get('faculty'):
                    dept_info.append(branding['faculty'])
                if branding.get('department'):
                    dept_info.append(branding['department'])
                
                if dept_info:
                    dept_style = ParagraphStyle(
                        name='DepartmentInfo',
                        parent=self.styles['Normal'],
                        fontSize=12,
                        alignment=TA_CENTER,
                        spaceAfter=5,
                        textColor=colors.black,
                        fontName=getattr(self, 'unicode_font_normal', 'Helvetica')
                    )
                    elements.append(Paragraph(' | '.join(dept_info), dept_style))
                
                if branding.get('course'):
                    course_style = ParagraphStyle(
                        name='CourseInfo',
                        parent=self.styles['Normal'],
                        fontSize=14,
                        alignment=TA_CENTER,
                        spaceAfter=15,
                        textColor=colors.black,
                        fontName=getattr(self, 'unicode_font_bold', 'Helvetica-Bold')
                    )
                    elements.append(Paragraph(branding['course'], course_style))
        
        # Exam Title - Very prominent with enhanced styling
        title = quiz_data.get('title', 'Quiz')
        title_style = ParagraphStyle(
            name='ExamTitle',
            parent=self.styles['Title'],
            fontSize=28,  # Increased from 24
            spaceAfter=25,
            spaceBefore=40,
            alignment=TA_CENTER,
            textColor=colors.darkblue,  # Changed to dark blue for distinction
            fontName=getattr(self, 'unicode_font_bold', 'Helvetica-Bold'),
            borderWidth=3,
            borderColor=colors.darkblue,
            borderPadding=20,
            backColor=colors.lightgrey
        )
        elements.append(Paragraph(title.upper(), title_style))
        
        # Exam information box - comprehensive
        info_parts = []
        if quiz_data.get('estimated_duration'):
            info_parts.append(f"Duration: {quiz_data['estimated_duration']}")
        if quiz_data.get('total_points'):
            info_parts.append(f"Total Points: {quiz_data['total_points']}")
        else:
            info_parts.append(f"Total Questions: {len(quiz_data.get('questions', []))}")
        
        if branding:
            if branding.get('instructor'):
                info_parts.append(f"Instructor: {branding['instructor']}")
            if branding.get('exam_date'):
                info_parts.append(f"Date: {branding['exam_date']}")
        
        if info_parts:
            info_style = ParagraphStyle(
                name='ExamInfo',
                parent=self.styles['Normal'],
                fontSize=13,  # Increased from 12
                alignment=TA_CENTER,
                spaceBefore=20,
                spaceAfter=25,
                leftIndent=30,
                rightIndent=30,
                fontName=getattr(self, 'unicode_font_bold', 'Helvetica-Bold'),
                borderWidth=2,
                borderColor=colors.darkblue,
                borderPadding=20,
                backColor=colors.beige,
                leading=18  # Better line spacing
            )
            info_text = "<br/>".join(info_parts)
            elements.append(Paragraph(info_text, info_style))
        
        # Instructions - compact but complete
        instructions_text = '''
        <b>INSTRUCTIONS:</b><br/><br/>
        • Read each question carefully and completely before answering<br/>
        • For multiple choice questions, select the best answer<br/>
        • Write clearly and legibly for all written responses<br/>
        • Show all work for calculation problems where applicable<br/>
        • Review your answers before submitting<br/>
        • Ask the instructor if you have any questions
        '''
        
        inst_style = ParagraphStyle(
            name='Instructions',
            parent=self.styles['Normal'],
            fontSize=12,  # Increased from 11
            spaceBefore=25,
            spaceAfter=25,
            leftIndent=25,
            rightIndent=25,
            fontName=getattr(self, 'unicode_font_normal', 'Helvetica'),
            borderWidth=2,
            borderColor=colors.darkgreen,
            borderPadding=18,
            backColor=colors.lightgreen,
            leading=16
        )
        elements.append(Paragraph(instructions_text, inst_style))
        
        # Student information section - enhanced with better spacing
        elements.append(Spacer(1, 40))
        
        # Add "Student Information" header
        student_header_style = ParagraphStyle(
            name='StudentHeader',
            parent=self.styles['Normal'],
            fontSize=14,
            alignment=TA_CENTER,
            spaceBefore=10,
            spaceAfter=15,
            fontName=getattr(self, 'unicode_font_bold', 'Helvetica-Bold'),
            textColor=colors.darkblue
        )
        elements.append(Paragraph('STUDENT INFORMATION', student_header_style))
        
        # Create student information table - more compact
        student_fields = []
        if branding and 'student_info' in branding:
            student_info = branding['student_info']
            if student_info.get('include_student_name', True):
                student_fields.append(['Student Full Name:', '_' * 60])
            if student_info.get('include_student_id', True):
                student_fields.append(['Student ID/Number:', '_' * 35])
            if student_info.get('include_date_field', False):
                student_fields.append(['Date:', '_' * 25])
            if student_info.get('include_signature', True):
                student_fields.append(['Signature:', '_' * 50])
        else:
            # Default comprehensive student fields
            student_fields = [
                ['Student Full Name:', '_' * 60],
                ['Student ID/Number:', '_' * 35],
                ['Signature:', '_' * 50]
            ]
        
        if student_fields:
            student_table = Table(student_fields, colWidths=[170, 380])
            student_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 12),  # Increased from 11
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 15),  # Increased padding
                ('TOPPADDING', (0, 0), (-1, -1), 10),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 1, colors.darkblue),  # Add grid
                ('BACKGROUND', (0, 0), (0, -1), colors.lightblue),  # Label background
            ]))
            elements.append(student_table)
        
        return elements
    
    def _create_cover_page(self, quiz_data: Dict[str, Any], branding: Dict[str, Any] = None) -> List:
        """Legacy method - kept for compatibility"""
        return self._create_single_cover_page(quiz_data, branding)
        
        # University Header - more prominent for cover page
        if branding:
            elements.extend(self._add_professional_branding(branding))
        
        # Large title - centered and prominent
        title = quiz_data.get('title', 'Quiz')
        cover_title_style = ParagraphStyle(
            name='CoverTitle',
            parent=self.styles['Title'],
            fontSize=32,
            spaceAfter=30,
            spaceBefore=80,
            alignment=TA_CENTER,
            textColor=colors.black,
            fontName=getattr(self, 'unicode_font_bold', 'Helvetica-Bold')
        )
        elements.append(Paragraph(title.upper(), cover_title_style))
        
        # Course and exam information box
        info_style = ParagraphStyle(
            name='CoverInfo',
            parent=self.styles['Normal'],
            fontSize=14,
            alignment=TA_CENTER,
            spaceBefore=30,
            spaceAfter=30,
            leftIndent=50,
            rightIndent=50,
            fontName=getattr(self, 'unicode_font_normal', 'Helvetica'),
            borderWidth=2,
            borderColor=colors.black,
            borderPadding=20,
            backColor=colors.lightgrey
        )
        
        # Build information text
        info_parts = []
        if quiz_data.get('estimated_duration'):
            info_parts.append(f"<b>Duration:</b> {quiz_data['estimated_duration']}")
        if quiz_data.get('total_points'):
            info_parts.append(f"<b>Total Points:</b> {quiz_data['total_points']}")
        else:
            info_parts.append(f"<b>Total Questions:</b> {len(quiz_data.get('questions', []))}")
        
        if branding:
            if branding.get('instructor'):
                info_parts.append(f"<b>Instructor:</b> {branding['instructor']}")
            if branding.get('exam_date'):
                info_parts.append(f"<b>Date:</b> {branding['exam_date']}")
        
        if info_parts:
            info_text = "<br/>".join(info_parts)
            elements.append(Paragraph(info_text, info_style))
        
        # Instructions preview on cover page
        cover_instructions = '''
        <b>EXAMINATION INSTRUCTIONS:</b><br/><br/>
        • Read all instructions carefully before beginning<br/>
        • Write your answers clearly in the spaces provided<br/>
        • For multiple choice questions, select the best answer<br/>
        • Show all work where applicable<br/>
        • Review your answers before submitting<br/>
        • Ask the instructor if you have any questions
        '''
        
        cover_inst_style = ParagraphStyle(
            name='CoverInstructions',
            parent=self.styles['Normal'],
            fontSize=12,
            spaceBefore=50,
            spaceAfter=30,
            leftIndent=40,
            rightIndent=40,
            fontName=getattr(self, 'unicode_font_normal', 'Helvetica'),
            borderWidth=1,
            borderColor=colors.grey,
            borderPadding=15,
            backColor=colors.beige
        )
        
        elements.append(Paragraph(cover_instructions, cover_inst_style))
        
        # Student information section on cover page
        elements.append(Spacer(1, 40))
        
        student_info_style = ParagraphStyle(
            name='CoverStudentInfo',
            parent=self.styles['Normal'],
            fontSize=14,
            alignment=TA_LEFT,
            fontName=getattr(self, 'unicode_font_normal', 'Helvetica')
        )
        
        # Create student information fields
        student_fields = []
        if branding and 'student_info' in branding:
            student_info = branding['student_info']
            if student_info.get('include_student_name', True):
                student_fields.append(['<b>Student Full Name:</b>', '_' * 50])
            if student_info.get('include_student_id', True):
                student_fields.append(['<b>Student ID/Number:</b>', '_' * 30])
            if student_info.get('include_date_field', False):
                student_fields.append(['<b>Date:</b>', '_' * 20])
            if student_info.get('include_signature', True):
                student_fields.append(['<b>Signature:</b>', '_' * 40])
        else:
            # Default student fields
            student_fields = [
                ['<b>Student Full Name:</b>', '_' * 50],
                ['<b>Student ID/Number:</b>', '_' * 30],
                ['<b>Signature:</b>', '_' * 40]
            ]
        
        if student_fields:
            student_table = Table(student_fields, colWidths=[160, 350])
            student_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 12),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 15),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            elements.append(student_table)
        
        return elements
    
    def _add_professional_branding(self, branding: Dict[str, Any]) -> List:
        """Add comprehensive professional university branding to document"""
        elements = []
        
        # University header with proper logo support
        if branding.get('university_name') or branding.get('institution_name'):
            institution = branding.get('university_name') or branding.get('institution_name')
            
            # Add logo if provided
            if branding.get('logo_path') or branding.get('has_logo'):
                try:
                    # Try to load actual logo image - check multiple path sources
                    logo_path = None
                    
                    # Check different logo path sources - try multiple methods
                    logo_path = None
                    
                    # Method 1: Direct path from branding settings
                    if branding.get('logo_path'):
                        potential_path = branding['logo_path']
                        if os.path.exists(potential_path):
                            logo_path = potential_path
                            logger.info(f"Found logo using direct path: {logo_path}")
                    
                    # Method 2: Convert URL to file path if direct path didn't work
                    if not logo_path and branding.get('logo_url'):
                        from django.conf import settings
                        logo_url = branding['logo_url']
                        if logo_url.startswith('/media/'):
                            # Remove /media/ prefix and join with MEDIA_ROOT
                            relative_path = logo_url.replace('/media/', '').lstrip('/')
                            potential_path = os.path.join(settings.MEDIA_ROOT, relative_path)
                            if os.path.exists(potential_path):
                                logo_path = potential_path
                                logger.info(f"Found logo using URL conversion: {logo_path}")
                    
                    # Method 3: Try using filename from branding settings
                    if not logo_path and branding.get('logo_filename'):
                        from django.conf import settings
                        filename = branding['logo_filename']
                        # Try export_logos directory
                        potential_path = os.path.join(settings.MEDIA_ROOT, 'export_logos', os.path.basename(filename))
                        if os.path.exists(potential_path):
                            logo_path = potential_path
                            logger.info(f"Found logo using filename: {logo_path}")
                    
                    if logo_path and os.path.exists(logo_path):
                        # Load and display actual logo image
                        from reportlab.platypus import Image as ReportLabImage
                        from reportlab.lib.utils import ImageReader
                        
                        # Create logo image with proper sizing and positioning
                        try:
                            # Center the logo
                            logo_style = ParagraphStyle(
                                name='LogoContainer',
                                parent=self.styles['Normal'],
                                alignment=TA_CENTER,
                                spaceAfter=15
                            )
                            
                            # Create image with appropriate size
                            logo_img = ReportLabImage(logo_path, width=80, height=80)
                            
                            # Create a table to center the logo properly
                            logo_table = Table([[logo_img]], colWidths=[6*inch])
                            logo_table.setStyle(TableStyle([
                                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                            ]))
                            
                            elements.append(logo_table)
                            elements.append(Spacer(1, 15))
                            logger.info(f"Successfully added university logo from: {logo_path}")
                        except Exception as img_error:
                            logger.error(f"Error creating logo image: {img_error}")
                            # Fallback to placeholder
                            logo_placeholder = Paragraph(
                                '''<para align="center">
                                <font size="12" color="gray">[UNIVERSITY LOGO]</font><br/>
                                <font size="8" color="gray">Image loading error</font>
                                </para>''',
                                self.styles['Normal']
                            )
                            elements.append(logo_placeholder)
                            elements.append(Spacer(1, 15))
                    else:
                        # Fallback to placeholder if logo file not found
                        logo_placeholder = Paragraph(
                            '''<para align="center">
                            <font size="12" color="gray">[UNIVERSITY LOGO]</font><br/>
                            <font size="8" color="gray">Upload logo in export form</font>
                            </para>''',
                            self.styles['Normal']
                        )
                        elements.append(logo_placeholder)
                        elements.append(Spacer(1, 15))
                        logger.info("Logo placeholder added - no valid logo path found")
                except Exception as e:
                    logger.error(f"Could not add logo: {e}")
                    # Fallback to text placeholder
                    logo_placeholder = Paragraph(
                        '''<para align="center">
                        <font size="12" color="gray">[UNIVERSITY LOGO]</font><br/>
                        <font size="8" color="gray">Logo processing error</font>
                        </para>''',
                        self.styles['Normal']
                    )
                    elements.append(logo_placeholder)
                    elements.append(Spacer(1, 15))
            
            # Comprehensive university info with full hierarchy
            uni_info = f"<b>{institution.upper()}</b>"
            
            # Add faculty if provided
            if branding.get('faculty'):
                uni_info += f"<br/><i>{branding['faculty']}</i>"
                
            # Add department
            if branding.get('department'):
                uni_info += f"<br/>{branding['department']}"
                
            # Add course information
            if branding.get('course'):
                uni_info += f"<br/><b>{branding['course']}</b>"
                
            # Add academic year and semester
            academic_info = []
            if branding.get('academic_year'):
                academic_info.append(f"Academic Year: {branding['academic_year']}")
            if branding.get('semester'):
                academic_info.append(branding['semester'])
            
            if academic_info:
                uni_info += f"<br/><i>{' | '.join(academic_info)}</i>"
            
            # Create a cleaner header layout without the problematic logo placeholder
            # University info in a clean, professional format
            header_elements = []
            
            # University name - main header
            university_style = ParagraphStyle(
                name='UniversityHeader',
                parent=self.styles['Normal'],
                fontSize=18,
                alignment=TA_CENTER,
                spaceBefore=0,
                spaceAfter=10,
                textColor=colors.black,
                fontName=getattr(self, 'unicode_font_bold', 'Helvetica-Bold')
            )
            header_elements.append(Paragraph(institution.upper(), university_style))
            
            # Faculty and department info
            faculty_info = []
            if branding.get('faculty'):
                faculty_info.append(branding['faculty'])
            if branding.get('department'):
                faculty_info.append(branding['department'])
            
            if faculty_info:
                faculty_style = ParagraphStyle(
                    name='FacultyHeader',
                    parent=self.styles['Normal'],
                    fontSize=12,
                    alignment=TA_CENTER,
                    spaceAfter=8,
                    textColor=colors.black,
                    fontName=getattr(self, 'unicode_font_normal', 'Helvetica')
                )
                faculty_text = ' | '.join(faculty_info)
                header_elements.append(Paragraph(faculty_text, faculty_style))
            
            # Course information
            if branding.get('course'):
                course_style = ParagraphStyle(
                    name='CourseHeader',
                    parent=self.styles['Normal'],
                    fontSize=14,
                    alignment=TA_CENTER,
                    spaceAfter=15,
                    textColor=colors.black,
                    fontName=getattr(self, 'unicode_font_bold', 'Helvetica-Bold')
                )
                header_elements.append(Paragraph(branding['course'], course_style))
            
            # Add all header elements
            for element in header_elements:
                elements.append(element)
            
            # Add a separator line
            elements.append(Spacer(1, 10))
            # Header elements already added above
            elements.append(Spacer(1, 20))
            
            # Enhanced exam metadata with instructor and date
            metadata_parts = []
            if branding.get('instructor'):
                metadata_parts.append(f"Instructor: {branding['instructor']}")
            if branding.get('exam_date'):
                metadata_parts.append(f"Date: {branding['exam_date']}")
            
            if metadata_parts:
                metadata_text = ' | '.join(metadata_parts)
                metadata_style = ParagraphStyle(
                    name='ExamMetadata',
                    parent=self.styles['Normal'],
                    fontSize=10,
                    alignment=TA_CENTER,
                    spaceAfter=15,
                    fontName=getattr(self, 'unicode_font_normal', 'Helvetica'),
                    borderWidth=1,
                    borderColor=colors.lightgrey,
                    borderPadding=8
                )
                elements.append(Paragraph(metadata_text, metadata_style))
        
        return elements
    
    def export_exam(self, exam_data: Dict[str, Any], branding: Dict[str, Any] = None) -> io.BytesIO:
        """Export exam to PDF format"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        # Add header/branding
        if branding:
            story.extend(self._add_branding(branding))
        
        # Title
        title = exam_data.get('title', 'Comprehensive Exam')
        story.append(Paragraph(title, self.styles['CustomTitle']))
        story.append(Spacer(1, 20))
        
        # Exam info
        info_data = [
            ['Duration:', f"{exam_data.get('duration', 120)} minutes"],
            ['Total Questions:', str(exam_data.get('total_questions', 0))],
            ['Date:', datetime.now().strftime('%B %d, %Y')],
            ['Time:', '________________'],
            ['Student Name:', '_' * 30],
            ['Student ID:', '_' * 20]
        ]
        
        info_table = Table(info_data, colWidths=[2*inch, 3*inch])
        info_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 20))
        
        # Instructions
        story.append(Paragraph("<b>INSTRUCTIONS:</b>", self.styles['CustomHeader']))
        instructions = """
        1. Read all instructions carefully before beginning<br/>
        2. Write your answers clearly in the spaces provided<br/>
        3. For multiple choice questions, circle the letter of your answer<br/>
        4. Show all work for calculation problems<br/>
        5. Check your answers before submitting
        """
        story.append(Paragraph(instructions, self.styles['Normal']))
        story.append(PageBreak())
        
        # Exam sections
        question_num = 1
        for section in exam_data.get('sections', []):
            # Section header
            section_title = f"Section {len(story)//20 + 1}: {section.get('name', 'Questions')}"
            story.append(Paragraph(section_title, self.styles['CustomHeader']))
            
            if section.get('instructions'):
                story.append(Paragraph(section['instructions'], self.styles['Normal']))
                story.append(Spacer(1, 10))
            
            # Section questions
            for question in section.get('questions', []):
                # Question text
                q_text = f"<b>{question_num}. {question.get('question', '')}</b> ({question.get('points', 1)} points)"
                story.append(Paragraph(q_text, self.styles['Question']))
                
                # Question options or answer space - clean professional format
                question_type = question.get('type', 'multiple_choice')
                if question_type == 'multiple_choice' and question.get('options'):
                    for j, option in enumerate(question['options']):
                        option_letter = chr(65 + j)
                        story.append(Paragraph(f"{option_letter}. {option}", self.styles['Option']))
                elif question_type == 'true_false':
                    story.append(Paragraph("A. True", self.styles['Option']))
                    story.append(Paragraph("B. False", self.styles['Option']))
                else:
                    # Add answer space for other question types
                    story.append(Spacer(1, 8))
                    for _ in range(4):  # 4 lines for answer
                        story.append(Paragraph("_" * 80, self.styles['Normal']))
                        story.append(Spacer(1, 6))
                
                story.append(Spacer(1, 15))
                question_num += 1
            
            story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def export_html_to_pdf(self, html_content: str, branding: Dict[str, Any] = None) -> io.BytesIO:
        """Convert HTML to PDF using ReportLab (fallback for WeasyPrint)"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        # Add branding if available
        if branding:
            story.extend(self._add_branding(branding))
        
        # Simple HTML to ReportLab conversion
        # This is a basic implementation - for complex HTML, WeasyPrint would be preferred
        from html.parser import HTMLParser
        
        class SimpleHTMLParser(HTMLParser):
            def __init__(self):
                super().__init__()
                self.current_text = ""
                self.in_title = False
                self.in_question = False
                
            def handle_starttag(self, tag, attrs):
                if tag == 'h1' or tag == 'h2':
                    self.in_title = True
                elif tag == 'div' and any(cls for name, cls in attrs if name == 'class' and 'question' in cls):
                    self.in_question = True
                    
            def handle_endtag(self, tag):
                if tag == 'h1' or tag == 'h2':
                    if self.current_text.strip():
                        story.append(Paragraph(self.current_text.strip(), self.styles['CustomTitle']))
                        self.current_text = ""
                    self.in_title = False
                elif tag == 'div' and self.in_question:
                    if self.current_text.strip():
                        story.append(Paragraph(self.current_text.strip(), self.styles['Question']))
                        self.current_text = ""
                    self.in_question = False
                    
            def handle_data(self, data):
                self.current_text += data
        
        # Parse HTML and convert to PDF
        parser = SimpleHTMLParser()
        parser.styles = self.styles
        
        try:
            parser.feed(html_content)
            
            # Add any remaining content
            if parser.current_text.strip():
                story.append(Paragraph(parser.current_text.strip(), self.styles['Normal']))
                
        except Exception as e:
            # Fallback: just add the HTML as plain text
            import re
            clean_text = re.sub('<.*?>', '', html_content)
            story.append(Paragraph(clean_text, self.styles['Normal']))
        
        # Build PDF
        if not story:
            story.append(Paragraph("Document content could not be processed.", self.styles['Normal']))
            
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def export_answer_key(self, content_data: Dict[str, Any], branding: Dict[str, Any] = None) -> io.BytesIO:
        """Export answer key to PDF format"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        # Add header/branding
        if branding:
            story.extend(self._add_branding(branding))
        
        # Title
        title = f"{content_data.get('title', 'Quiz')} - Answer Key"
        story.append(Paragraph(title, self.styles['CustomTitle']))
        story.append(Spacer(1, 20))
        
        # Answers
        if 'questions' in content_data:
            questions = content_data['questions']
        else:
            # Handle exam format with sections
            questions = []
            for section in content_data.get('sections', []):
                questions.extend(section.get('questions', []))
        
        for i, question in enumerate(questions, 1):
            # Question and answer
            q_text = f"<b>{i}. {question.get('question', '')}</b>"
            story.append(Paragraph(q_text, self.styles['Question']))
            
            answer_text = f"<b>Answer:</b> {question.get('correct_answer', 'Not provided')}"
            story.append(Paragraph(answer_text, self.styles['Answer']))
            
            # Explanation if available
            if question.get('explanation'):
                exp_text = f"<b>Explanation:</b> {question['explanation']}"
                story.append(Paragraph(exp_text, self.styles['Normal']))
            
            story.append(Spacer(1, 10))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def _add_branding(self, branding: Dict[str, Any]) -> List:
        """Add professional branding elements to document"""
        elements = []
        branding = branding or {}

        # Logo (embedded) - keeps PDF self-contained and avoids path issues
        logo_flowable = self._build_logo_flowable(branding, max_w=1.0*inch, max_h=1.0*inch)
        if logo_flowable:
            logo_table = Table([[logo_flowable]], colWidths=[6*inch])
            logo_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            elements.append(logo_table)
            elements.append(Spacer(1, 10))
        
        # Institution name - university style
        if branding.get('institution_name') or branding.get('university_name'):
            institution = branding.get('institution_name') or branding.get('university_name')
            inst_style = ParagraphStyle(
                name='Institution',
                parent=self.styles['Normal'],
                fontSize=16,
                alignment=TA_CENTER,
                spaceBefore=0,
                spaceAfter=8,
                textColor=colors.black,
                fontName='Helvetica-Bold'
            )
            elements.append(Paragraph(institution.upper(), inst_style))
        
        # Department
        if branding.get('department'):
            dept_style = ParagraphStyle(
                name='Department',
                parent=self.styles['Normal'],
                fontSize=12,
                alignment=TA_CENTER,
                spaceAfter=6,
                textColor=colors.black,
                fontName='Helvetica'
            )
            elements.append(Paragraph(branding['department'], dept_style))
        
        # Course information
        if branding.get('course'):
            course_style = ParagraphStyle(
                name='Course',
                parent=self.styles['Normal'],
                fontSize=11,
                alignment=TA_CENTER,
                spaceAfter=20,
                textColor=colors.black,
                fontName='Helvetica'
            )
            course_text = branding['course']
            if branding.get('semester'):
                course_text += f" - {branding['semester']}"
            elements.append(Paragraph(course_text, course_style))
        
        return elements


# Define PDF_AVAILABLE alias to avoid conflicts
PDF_AVAILABLE = REPORTLAB_AVAILABLE

class DOCXExporter:
    """Service for exporting content to DOCX format"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX export")
    
    def export_quiz(self, quiz_data: Dict[str, Any], branding: Dict[str, Any] = None) -> io.BytesIO:
        """Export quiz to DOCX format with professional university formatting"""
        doc = Document()
        
        # Set document margins for professional look
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        
        # Add single comprehensive cover page content
        if branding:
            self._add_single_docx_cover_page(doc, quiz_data, branding)
        else:
            # Simple title if no branding
            title = doc.add_heading(quiz_data.get('title', 'Quiz'), level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a clear "QUESTIONS" section header after cover page
        questions_heading = doc.add_heading('QUESTIONS', level=1)
        questions_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Questions with professional formatting
        for i, question in enumerate(quiz_data.get('questions', []), 1):
            question_type = question.get('type', 'multiple_choice')
            points = question.get('points', 1)
            
            # Question header
            q_header = doc.add_paragraph()
            q_header.add_run(f'Question {i}.').bold = True
            q_header.add_run(f' ({points} point{"s" if points != 1 else ""})')
            
            # Question text with justified alignment
            q_text_para = doc.add_paragraph(question.get('question', ''))
            q_text_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Handle different question types
            if question_type == 'multiple_choice' and question.get('options'):
                doc.add_paragraph()  # Add space before options
                # Remove duplicate options and support up to 5 options (A-E)
                unique_options = list(dict.fromkeys(question['options']))[:5]
                for j, option in enumerate(unique_options):
                    option_letter = chr(65 + j)  # A, B, C, D, E
                    option_para = doc.add_paragraph(f'{option_letter}. {option}')
                    option_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    # Indent options slightly
                    option_para.paragraph_format.left_indent = Inches(0.3)
            
            elif question_type == 'true_false':
                doc.add_paragraph()
                true_para = doc.add_paragraph('A. True')
                true_para.paragraph_format.left_indent = Inches(0.3)
                false_para = doc.add_paragraph('B. False')
                false_para.paragraph_format.left_indent = Inches(0.3)
            
            elif question_type == 'short_answer':
                doc.add_paragraph()
                answer_label = doc.add_paragraph()
                answer_label.add_run('Answer:').bold = True
                doc.add_paragraph()
                # Create table for neat answer lines
                table = doc.add_table(rows=5, cols=1)
                for row in table.rows:
                    row.cells[0].text = ''
                    row.height = Inches(0.3)
            
            elif question_type == 'fill_blank':
                doc.add_paragraph()
                fill_para = doc.add_paragraph()
                fill_para.add_run('Answer: ').bold = True
                fill_para.add_run('_' * 50)
            
            elif question_type == 'essay':
                doc.add_paragraph()
                essay_label = doc.add_paragraph()
                essay_label.add_run('Answer: ').bold = True
                essay_label.add_run('(Use the space below for your complete response)')
                doc.add_paragraph()
                # Create larger table for essay responses
                essay_table = doc.add_table(rows=10, cols=1)
                for row in essay_table.rows:
                    row.cells[0].text = ''
                    row.height = Inches(0.35)
            
            # Add compact space between questions
            doc.add_paragraph()  # Single paragraph space instead of double
        
        # Add watermark if specified
        if branding and branding.get('watermark'):
            watermark_text = branding['watermark']
            if watermark_text.strip():
                try:
                    # Add watermark as a text box or header
                    # For DOCX, we'll add it as a subtle footer text
                    sections = doc.sections
                    for section in sections:
                        footer = section.footer
                        footer_para = footer.paragraphs[0]
                        footer_run = footer_para.add_run(f"  [{watermark_text}]  ")
                        footer_run.font.size = Pt(8)
                        footer_run.font.color.rgb = RGBColor(128, 128, 128)  # Light gray
                        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    logger.warning(f"Could not add watermark to DOCX: {e}")
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    
    def _add_single_docx_cover_page(self, doc: Document, quiz_data: Dict[str, Any], branding: Dict[str, Any]):
        """Add single comprehensive cover page to DOCX document"""
        # Add logo first if available (bytes-based so it works with non-local storage)
        logo_resolved = _read_logo_bytes_from_branding(branding)
        if logo_resolved:
            try:
                logo_bytes, mime = logo_resolved
                logo_bytes, _ = _prepare_logo_for_embedding(logo_bytes, mime, max_px=256)

                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_run = logo_para.add_run()
                logo_run.add_picture(io.BytesIO(logo_bytes), width=Inches(1.5))
                doc.add_paragraph()  # spacing after logo
                logger.info("Successfully added logo to DOCX cover page")
            except Exception as e:
                logger.warning(f"Could not add logo to DOCX cover page: {e}")
        
        # University information
        if branding.get('university_name') or branding.get('institution_name'):
            institution = branding.get('university_name') or branding.get('institution_name')
            uni_para = doc.add_paragraph(institution.upper())
            uni_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            uni_para.runs[0].font.size = Pt(18)
            uni_para.runs[0].bold = True
            
            # Department and course info
            dept_info = []
            if branding.get('faculty'):
                dept_info.append(branding['faculty'])
            if branding.get('department'):
                dept_info.append(branding['department'])
            
            if dept_info:
                dept_para = doc.add_paragraph(' | '.join(dept_info))
                dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                dept_para.runs[0].font.size = Pt(12)
            
            if branding.get('course'):
                course_para = doc.add_paragraph(branding['course'])
                course_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                course_para.runs[0].font.size = Pt(14)
                course_para.runs[0].bold = True
        
        # Exam title
        title = doc.add_heading(quiz_data.get('title', 'Quiz').upper(), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Exam info
        info_parts = []
        if quiz_data.get('estimated_duration'):
            info_parts.append(f"Duration: {quiz_data['estimated_duration']}")
        if quiz_data.get('total_points'):
            info_parts.append(f"Total Points: {quiz_data['total_points']}")
        else:
            info_parts.append(f"Total Questions: {len(quiz_data.get('questions', []))}")
        
        if branding:
            if branding.get('instructor'):
                info_parts.append(f"Instructor: {branding['instructor']}")
            if branding.get('exam_date'):
                info_parts.append(f"Date: {branding['exam_date']}")
        
        if info_parts:
            info_para = doc.add_paragraph(' | '.join(info_parts))
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_para.runs[0].italic = True
        
        # Instructions
        doc.add_heading('INSTRUCTIONS:', level=2)
        instructions = doc.add_paragraph()
        instructions.add_run('• Read each question carefully and completely before answering\n')
        instructions.add_run('• For multiple choice questions, select the best answer\n')
        instructions.add_run('• Write clearly and legibly for all written responses\n')
        instructions.add_run('• Show all work for calculation problems where applicable\n')
        instructions.add_run('• Review your answers before submitting\n')
        instructions.add_run('• Ask the instructor if you have any questions')
        
        # Student information
        student_info = branding.get('student_info', {}) if branding else {}
        if any(student_info.get(field, True) for field in ['include_student_name', 'include_student_id', 'include_signature', 'include_date_field']):
            doc.add_paragraph()  # Space
            
            student_fields = []
            if student_info.get('include_student_name', True):
                student_fields.append(['Student Full Name:', '_' * 50])
            if student_info.get('include_student_id', True):
                student_fields.append(['Student ID/Number:', '_' * 30])
            if student_info.get('include_date_field', False):
                student_fields.append(['Date:', '_' * 20])
            if student_info.get('include_signature', True):
                student_fields.append(['Signature:', '_' * 40])
            
            if student_fields:
                student_table = doc.add_table(rows=len(student_fields), cols=2)
                student_table.columns[0].width = Inches(2.0)
                student_table.columns[1].width = Inches(4.0)
                
                for i, (label, line) in enumerate(student_fields):
                    label_cell = student_table.cell(i, 0)
                    line_cell = student_table.cell(i, 1)
                    
                    label_para = label_cell.paragraphs[0]
                    label_para.text = label
                    label_para.runs[0].bold = True
                    
                    line_para = line_cell.paragraphs[0]
                    line_para.text = line
        
        # Page break after cover page
        doc.add_page_break()
    
    def _add_professional_docx_branding(self, doc: Document, branding: Dict[str, Any]):
        """Legacy method - kept for compatibility"""
        # Enhanced university header table
        if branding.get('university_name') or branding.get('institution_name'):
            institution = branding.get('university_name') or branding.get('institution_name')
            
            # Create header table
            header_table = doc.add_table(rows=1, cols=2)
            header_table.columns[0].width = Inches(1.8)
            header_table.columns[1].width = Inches(4.7)
            
            # Enhanced logo cell
            logo_cell = header_table.cell(0, 0)
            logo_para = logo_cell.paragraphs[0]
            
            # Try to add actual logo image
            logo_added = False
            logo_path = None
            
            # Check different logo path sources - try multiple methods
            logo_path = None
            
            # Method 1: Direct path from branding settings
            if branding.get('logo_path'):
                potential_path = branding['logo_path']
                if os.path.exists(potential_path):
                    logo_path = potential_path
                    logger.info(f"Found logo for DOCX using direct path: {logo_path}")
            
            # Method 2: Convert URL to file path if direct path didn't work
            if not logo_path and branding.get('logo_url'):
                from django.conf import settings
                logo_url = branding['logo_url']
                if logo_url.startswith('/media/'):
                    # Remove /media/ prefix and join with MEDIA_ROOT
                    relative_path = logo_url.replace('/media/', '').lstrip('/')
                    potential_path = os.path.join(settings.MEDIA_ROOT, relative_path)
                    if os.path.exists(potential_path):
                        logo_path = potential_path
                        logger.info(f"Found logo for DOCX using URL conversion: {logo_path}")
            
            # Method 3: Try using filename from branding settings
            if not logo_path and branding.get('logo_filename'):
                from django.conf import settings
                filename = branding['logo_filename']
                # Try export_logos directory
                potential_path = os.path.join(settings.MEDIA_ROOT, 'export_logos', os.path.basename(filename))
                if os.path.exists(potential_path):
                    logo_path = potential_path
                    logger.info(f"Found logo for DOCX using filename: {logo_path}")
                        
            if logo_path and os.path.exists(logo_path):
                try:
                    # Clear the placeholder text
                    logo_para.clear()
                    # Add the actual logo image
                    run = logo_para.add_run()
                    run.add_picture(logo_path, width=Inches(1.2))
                    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    logo_added = True
                    logger.info(f"Successfully added logo to DOCX: {logo_path}")
                except Exception as e:
                    logger.error(f"Could not add logo to DOCX: {e}")
            
            # Fallback to placeholder text if logo wasn't added
            if not logo_added:
                logo_para.text = '[UNIVERSITY LOGO]\nUpload logo in\nexport form'
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_para.runs[0].font.size = Pt(9)
            
            # Comprehensive university info cell
            info_cell = header_table.cell(0, 1)
            info_para = info_cell.paragraphs[0]
            
            # Add university name
            info_run = info_para.add_run(institution.upper())
            info_run.bold = True
            info_run.font.size = Pt(16)
            
            # Add faculty if provided
            if branding.get('faculty'):
                faculty_run = info_para.add_run('\n' + branding['faculty'])
                faculty_run.italic = True
                faculty_run.font.size = Pt(12)
            
            # Add department
            if branding.get('department'):
                dept_run = info_para.add_run('\n' + branding['department'])
                dept_run.font.size = Pt(11)
            
            # Add course information
            if branding.get('course'):
                course_run = info_para.add_run('\n' + branding['course'])
                course_run.bold = True
                course_run.font.size = Pt(12)
            
            # Add academic year and semester
            academic_info = []
            if branding.get('academic_year'):
                academic_info.append(f"Academic Year: {branding['academic_year']}")
            if branding.get('semester'):
                academic_info.append(branding['semester'])
            
            if academic_info:
                academic_run = info_para.add_run('\n' + ' | '.join(academic_info))
                academic_run.italic = True
                academic_run.font.size = Pt(10)
            
            # Style the table
            header_table.style = 'Table Grid'
            
            doc.add_paragraph()  # Space after header
            
            # Enhanced exam metadata
            if branding.get('instructor') or branding.get('exam_date'):
                metadata_parts = []
                if branding.get('instructor'):
                    metadata_parts.append(f'Instructor: {branding["instructor"]}')
                if branding.get('exam_date'):
                    metadata_parts.append(f'Date: {branding["exam_date"]}')
                
                metadata_para = doc.add_paragraph(' | '.join(metadata_parts))
                metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                metadata_para.runs[0].font.size = Pt(11)
                
                doc.add_paragraph()  # Space after metadata
        
        # Add student information section
        student_info = branding.get('student_info', {}) if branding else {}
        if any(student_info.get(field, True) for field in ['include_student_name', 'include_student_id', 'include_signature', 'include_date_field']):
            # Student information header
            student_heading = doc.add_heading('STUDENT INFORMATION', level=3)
            student_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Create student information table
            student_fields = []
            if student_info.get('include_student_name', True):
                student_fields.append(['Student Full Name:', '_' * 40])
            if student_info.get('include_student_id', True):
                student_fields.append(['Student ID/Number:', '_' * 25])
            if student_info.get('include_date_field', False):
                student_fields.append(['Date:', '_' * 15])
            if student_info.get('include_signature', True):
                student_fields.append(['Signature:', '_' * 30])
            
            if student_fields:
                student_table = doc.add_table(rows=len(student_fields), cols=2)
                student_table.columns[0].width = Inches(2.0)
                student_table.columns[1].width = Inches(4.0)
                
                for i, (label, line) in enumerate(student_fields):
                    label_cell = student_table.cell(i, 0)
                    line_cell = student_table.cell(i, 1)
                    
                    label_para = label_cell.paragraphs[0]
                    label_para.text = label
                    label_para.runs[0].bold = True
                    
                    line_para = line_cell.paragraphs[0]
                    line_para.text = line
                
                # Style the student table
                student_table.style = 'Table Grid'
                
                doc.add_paragraph()  # Space after student info
        
        # Instructions
        doc.add_heading('INSTRUCTIONS:', level=2)
        instructions = doc.add_paragraph()
        instructions.add_run('1. Read all instructions carefully before beginning\n')
        instructions.add_run('2. Write your answers clearly in the spaces provided\n')
        instructions.add_run('3. For multiple choice questions, circle the letter of your answer\n')
        instructions.add_run('4. Show all work for calculation problems\n')
        instructions.add_run('5. Check your answers before submitting')
        
    
    def _add_branding(self, doc: Document, branding: Dict[str, Any]):
        """Add professional branding to document"""
        # Institution name - university style
        if branding.get('institution_name') or branding.get('university_name'):
            institution = branding.get('institution_name') or branding.get('university_name')
            inst_para = doc.add_paragraph(institution.upper())
            inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            inst_para.runs[0].font.size = Pt(16)
            inst_para.runs[0].bold = True
        
        # Department
        if branding.get('department'):
            dept_para = doc.add_paragraph(branding['department'])
            dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dept_para.runs[0].font.size = Pt(12)
        
        # Course information
        if branding.get('course'):
            course_text = branding['course']
            if branding.get('semester'):
                course_text += f" - {branding['semester']}"
            course_para = doc.add_paragraph(course_text)
            course_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            course_para.runs[0].font.size = Pt(11)
        
        doc.add_paragraph()


class ZipExporter:
    """Service for creating ZIP archives of multiple export formats"""
    
    def __init__(self):
        self.pdf_exporter = PDFExporter() if REPORTLAB_AVAILABLE else None
        self.docx_exporter = DOCXExporter() if DOCX_AVAILABLE else None
        self.html_exporter = HTMLExporter()
    
    def create_multi_format_export(self, quiz_data: Dict[str, Any], branding: Dict[str, Any] = None,
                                 formats: List[str] = None, include_answer_key: bool = True) -> io.BytesIO:
        """Create a ZIP file containing the quiz/exam in multiple formats
        
        Args:
            quiz_data: The quiz/exam data
            branding: Branding information
            formats: List of formats to include ['pdf', 'docx', 'html']
            include_answer_key: Whether to include answer key versions
            
        Returns:
            BytesIO buffer containing the ZIP file
        """
        if formats is None:
            formats = ['pdf', 'html']
            if DOCX_AVAILABLE:
                formats.append('docx')
        
        # Clean quiz data by removing JSON/metadata that shouldn't be in exports
        cleaned_quiz_data = self._clean_quiz_data_for_export(quiz_data)
        
        zip_buffer = io.BytesIO()
        
        # Create safe filename base
        title = cleaned_quiz_data.get('title', 'Quiz')
        safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()
        safe_title = safe_title.replace(' ', '_')[:50]
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            
            # Student version exports
            if 'pdf' in formats and self.pdf_exporter:
                try:
                    pdf_buffer = self.pdf_exporter.export_quiz(cleaned_quiz_data, branding)
                    zip_file.writestr(f"{safe_title}_Student.pdf", pdf_buffer.getvalue())
                    pdf_buffer.close()
                except Exception as e:
                    logger.error(f"PDF export failed: {e}")
            
            if 'html' in formats:
                try:
                    html_content = self.html_exporter.export_quiz(cleaned_quiz_data, branding, show_answers=False)
                    zip_file.writestr(f"{safe_title}_Student.html", html_content.encode('utf-8'))
                except Exception as e:
                    logger.error(f"HTML export failed: {e}")
            
            if 'docx' in formats and self.docx_exporter:
                try:
                    docx_buffer = self.docx_exporter.export_quiz(cleaned_quiz_data, branding)
                    zip_file.writestr(f"{safe_title}_Student.docx", docx_buffer.getvalue())
                    docx_buffer.close()
                except Exception as e:
                    logger.error(f"DOCX export failed: {e}")
            
            # Answer key versions (instructor versions)
            if include_answer_key:
                if 'pdf' in formats and self.pdf_exporter:
                    try:
                        answer_key_buffer = self.pdf_exporter.export_answer_key(cleaned_quiz_data, branding)
                        zip_file.writestr(f"{safe_title}_Answer_Key.pdf", answer_key_buffer.getvalue())
                        answer_key_buffer.close()
                    except Exception as e:
                        logger.error(f"PDF answer key export failed: {e}")
                
                if 'html' in formats:
                    try:
                        instructor_html = self.html_exporter.export_quiz(cleaned_quiz_data, branding, show_answers=True)
                        zip_file.writestr(f"{safe_title}_Instructor.html", instructor_html.encode('utf-8'))
                    except Exception as e:
                        logger.error(f"HTML instructor version export failed: {e}")
            
            # Add export information file
            export_info = self._create_export_info(cleaned_quiz_data, branding, formats)
            zip_file.writestr("Export_Info.txt", export_info.encode('utf-8'))
        
        zip_buffer.seek(0)
        return zip_buffer
    
    def _clean_quiz_data_for_export(self, quiz_data: Dict[str, Any]) -> Dict[str, Any]:
        """Clean quiz data by removing internal metadata and JSON artifacts"""
        cleaned_data = quiz_data.copy() if quiz_data else {}
        
        # Remove internal/technical fields that shouldn't appear in exports
        fields_to_remove = [
            'metadata', 'tokens_used', 'processing_time', 'source_language',
            'target_language', 'fallback', 'api_response', 'raw_content',
            'generation_parameters', 'model_version', 'request_id'
        ]
        
        for field in fields_to_remove:
            cleaned_data.pop(field, None)
        
        # Clean questions data
        if 'questions' in cleaned_data:
            cleaned_questions = []
            for question in cleaned_data['questions']:
                if isinstance(question, dict):
                    cleaned_question = {
                        'id': question.get('id'),
                        'type': question.get('type'),
                        'question': question.get('question', ''),
                        'correct_answer': question.get('correct_answer', ''),
                        'explanation': question.get('explanation', ''),
                        'difficulty': question.get('difficulty', 'medium'),
                        'points': question.get('points', 1)
                    }
                    
                    # Add options for multiple choice questions (remove duplicates)
                    if question.get('type') == 'multiple_choice' and question.get('options'):
                        unique_options = list(dict.fromkeys(question['options']))  # Remove duplicates
                        cleaned_question['options'] = unique_options[:5]  # Support up to 5 options (A-E)
                    
                    cleaned_questions.append(cleaned_question)
            
            cleaned_data['questions'] = cleaned_questions
        
        # Clean sections data for exams
        if 'sections' in cleaned_data:
            cleaned_sections = []
            for section in cleaned_data['sections']:
                if isinstance(section, dict):
                    cleaned_section = {
                        'name': section.get('name', ''),
                        'instructions': section.get('instructions', ''),
                        'questions': [],
                        'points': section.get('points', 0)
                    }
                    
                    # Clean questions in section
                    for question in section.get('questions', []):
                        if isinstance(question, dict):
                            cleaned_question = {
                                'id': question.get('id'),
                                'type': question.get('type'),
                                'question': question.get('question', ''),
                                'correct_answer': question.get('correct_answer', ''),
                                'explanation': question.get('explanation', ''),
                                'difficulty': question.get('difficulty', 'medium'),
                                'points': question.get('points', 1)
                            }
                            
                            # Add options for multiple choice (remove duplicates)
                            if question.get('type') == 'multiple_choice' and question.get('options'):
                                unique_options = list(dict.fromkeys(question['options']))
                                cleaned_question['options'] = unique_options[:5]  # Support up to 5 options (A-E)
                            
                            cleaned_section['questions'].append(cleaned_question)
                    
                    cleaned_sections.append(cleaned_section)
            
            cleaned_data['sections'] = cleaned_sections
        
        return cleaned_data
    
    def _create_export_info(self, quiz_data: Dict[str, Any], branding: Dict[str, Any], formats: List[str]) -> str:
        """Create an information file about the export"""
        info_lines = [
            "=" * 50,
            "EXAM/QUIZ EXPORT INFORMATION",
            "=" * 50,
            "",
            f"Title: {quiz_data.get('title', 'Untitled')}",
            f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"Total Questions: {len(quiz_data.get('questions', []))}",
            ""
        ]
        
        if quiz_data.get('estimated_duration'):
            info_lines.append(f"Duration: {quiz_data['estimated_duration']}")
        
        if quiz_data.get('total_points'):
            info_lines.append(f"Total Points: {quiz_data['total_points']}")
        
        info_lines.extend([
            "",
            "EXPORTED FORMATS:",
            "-" * 20
        ])
        
        format_descriptions = {
            'pdf': 'PDF - Professional print-ready format with proper pagination',
            'html': 'HTML - Web-viewable format with responsive design',
            'docx': 'DOCX - Microsoft Word format for easy editing'
        }
        
        for fmt in formats:
            if fmt in format_descriptions:
                info_lines.append(f"• {format_descriptions[fmt]}")
        
        if branding:
            info_lines.extend([
                "",
                "INSTITUTION INFORMATION:",
                "-" * 25
            ])
            
            if branding.get('institution_name'):
                info_lines.append(f"Institution: {branding['institution_name']}")
            if branding.get('department'):
                info_lines.append(f"Department: {branding['department']}")
            if branding.get('instructor'):
                info_lines.append(f"Instructor: {branding['instructor']}")
        
        info_lines.extend([
            "",
            "FILES INCLUDED:",
            "-" * 15,
            "• Student version (for exam administration)",
            "• Answer key/Instructor version (with correct answers)",
            "• This information file",
            "",
            "Generated by DidactAI - Professional Educational Content Platform"
        ])
        
        return "\n".join(info_lines)


class HTMLExporter:
    """Service for exporting content to HTML format with professional university styling"""
    
    def export_quiz(self, quiz_data: Dict[str, Any], branding: Dict[str, Any] = None, show_answers: bool = False) -> str:
        """Export quiz to HTML format
        
        Args:
            quiz_data: The quiz data to export
            branding: Optional branding information
            show_answers: Whether to show correct answers (for instructor version)
        """
        context = {
            'quiz_data': quiz_data,
            'branding': branding or {},
            'export_date': datetime.now().strftime('%B %d, %Y')
        }
        
        html_template = '''
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{{ quiz_data.title }}</title>
            <style>
                @import url('https://fonts.googleapis.com/css2?family=Crimson+Text:wght@400;600;700&family=Inter:wght@300;400;500;600&display=swap');
                
                * {
                    margin: 0;
                    padding: 0;
                    box-sizing: border-box;
                }
                
                body {
                    font-family: 'Inter', 'Arial', sans-serif;
                    line-height: 1.6;
                    color: #2c3e50;
                    background: white;
                    margin: 0;
                    padding: 60px 40px 40px 40px;
                }
                
                .exam-container {
                    max-width: 210mm;
                    margin: 0 auto;
                    background: white;
                    min-height: 297mm;
                }
                
                /* University Header */
                .university-header {
                    display: flex;
                    align-items: center;
                    border-bottom: 3px solid #2c3e50;
                    padding-bottom: 20px;
                    margin-bottom: 30px;
                }
                
                .logo-placeholder {
                    width: 80px;
                    height: 80px;
                    border: 2px solid #34495e;
                    border-radius: 50%;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    font-size: 12px;
                    color: #7f8c8d;
                    margin-right: 25px;
                    background: #ecf0f1;
                    flex-shrink: 0;
                }
                
                .header-info {
                    flex: 1;
                }
                
                .university-name {
                    font-family: 'Crimson Text', serif;
                    font-size: 28px;
                    font-weight: 700;
                    color: #2c3e50;
                    margin-bottom: 5px;
                    text-transform: uppercase;
                    letter-spacing: 1px;
                }
                
                .faculty-name {
                    font-size: 18px;
                    color: #2c3e50;
                    margin-bottom: 6px;
                    font-weight: 600;
                    font-style: italic;
                }
                
                .department-name {
                    font-size: 16px;
                    color: #34495e;
                    margin-bottom: 4px;
                    font-weight: 500;
                }
                
                .course-info {
                    font-size: 14px;
                    color: #2c3e50;
                    font-weight: 600;
                    margin-bottom: 4px;
                }
                
                .academic-info {
                    font-size: 12px;
                    color: #7f8c8d;
                    font-weight: 400;
                    font-style: italic;
                }
                
                /* Exam Title */
                .exam-title {
                    text-align: center;
                    margin: 40px 0 30px;
                }
                
                .exam-type {
                    font-family: 'Crimson Text', serif;
                    font-size: 36px;
                    font-weight: 700;
                    color: #2c3e50;
                    margin-bottom: 10px;
                    text-transform: uppercase;
                    letter-spacing: 2px;
                }
                
                /* Exam Metadata */
                .exam-metadata {
                    display: grid;
                    grid-template-columns: 1fr 1fr;
                    gap: 15px;
                    margin: 30px 0;
                    padding: 20px;
                    border: 1px solid #bdc3c7;
                    background: #f8f9fa;
                }
                
                .metadata-item {
                    display: flex;
                    align-items: center;
                    font-size: 14px;
                }
                
                .metadata-label {
                    font-weight: 600;
                    color: #2c3e50;
                    margin-right: 10px;
                    min-width: 80px;
                }
                
                .metadata-value {
                    color: #34495e;
                    border-bottom: 1px solid #bdc3c7;
                    flex: 1;
                    padding-bottom: 2px;
                }
                
                /* Instructions Section */
                .instructions-section {
                    margin: 30px 0;
                    padding: 20px;
                    border: 2px solid #34495e;
                    background: #fdfdfd;
                }
                
                .instructions-title {
                    font-family: 'Crimson Text', serif;
                    font-size: 18px;
                    font-weight: 700;
                    color: #2c3e50;
                    margin-bottom: 15px;
                    text-transform: uppercase;
                    border-bottom: 1px solid #ecf0f1;
                    padding-bottom: 5px;
                }
                
                .instructions-list {
                    list-style: none;
                }
                
                .instructions-list li {
                    padding: 5px 0;
                    position: relative;
                    padding-left: 20px;
                    color: #34495e;
                    font-size: 14px;
                }
                
                .instructions-list li::before {
                    content: '-';
                    position: absolute;
                    left: 0;
                    color: #2c3e50;
                    font-weight: bold;
                }
                
                /* Questions */
                .question {
                    margin: 25px 0;
                    page-break-inside: avoid;
                }
                
                .question-header {
                    display: flex;
                    justify-content: space-between;
                    align-items: flex-start;
                    margin-bottom: 12px;
                }
                
                .question-number {
                    font-weight: 600;
                    color: #2c3e50;
                    font-size: 16px;
                }
                
                .question-points {
                    font-size: 12px;
                    color: #7f8c8d;
                    font-style: italic;
                }
                
                .question-text {
                    font-size: 15px;
                    line-height: 1.6;
                    color: #2c3e50;
                    margin-bottom: 15px;
                    text-align: justify;
                }
                
                /* Multiple Choice Options */
                .mc-options {
                    margin: 15px 0;
                }
                
                .mc-option {
                    display: flex;
                    margin: 8px 0;
                    align-items: flex-start;
                }
                
                .option-checkbox {
                    width: 12px;
                    height: 12px;
                    border: 2px solid #2c3e50;
                    margin-right: 10px;
                    margin-top: 4px;
                    flex-shrink: 0;
                }
                
                .option-letter {
                    font-weight: 600;
                    color: #2c3e50;
                    margin-right: 8px;
                    min-width: 20px;
                }
                
                .option-text {
                    line-height: 1.5;
                    color: #34495e;
                }
                
                /* True/False Questions */
                .tf-options {
                    display: flex;
                    gap: 40px;
                    margin: 15px 0;
                    justify-content: flex-start;
                }
                
                .tf-option {
                    display: flex;
                    align-items: center;
                }
                
                /* Answer Spaces */
                .answer-space {
                    border-bottom: 2px solid #2c3e50;
                    display: inline-block;
                    min-width: 300px;
                    height: 25px;
                    margin: 10px 0;
                }
                
                .answer-lines {
                    margin: 15px 0;
                }
                
                .answer-line {
                    border-bottom: 1px solid #95a5a6;
                    height: 30px;
                    margin: 8px 0;
                    width: 100%;
                }
                
                /* Fill in the Blank */
                .fill-blank {
                    margin: 15px 0;
                }
                
                .blank-space {
                    border-bottom: 2px solid #2c3e50;
                    display: inline-block;
                    min-width: 120px;
                    height: 20px;
                    margin: 0 5px;
                }
                
                /* Essay Questions */
                .essay-space {
                    margin: 20px 0;
                    min-height: 200px;
                    border: 1px solid #bdc3c7;
                    background: repeating-linear-gradient(
                        transparent,
                        transparent 28px,
                        #ecf0f1 28px,
                        #ecf0f1 30px
                    );
                }
                
                /* Student Information Section */
                .student-info-section {
                    margin: 30px 0;
                    padding: 20px;
                    border: 2px solid #34495e;
                    background: #fdfdfd;
                }
                
                .student-info-title {
                    font-family: 'Crimson Text', serif;
                    font-size: 16px;
                    font-weight: 700;
                    color: #2c3e50;
                    margin-bottom: 15px;
                    text-transform: uppercase;
                    border-bottom: 1px solid #ecf0f1;
                    padding-bottom: 5px;
                }
                
                .student-fields {
                    display: grid;
                    grid-template-columns: 1fr 1fr;
                    gap: 15px;
                }
                
                .student-field {
                    display: flex;
                    align-items: center;
                    font-size: 14px;
                }
                
                .student-field-label {
                    font-weight: 600;
                    color: #2c3e50;
                    margin-right: 10px;
                    min-width: 120px;
                }
                
                .student-field-line {
                    flex: 1;
                    border-bottom: 2px solid #2c3e50;
                    height: 25px;
                }
                
                .signature-field {
                    grid-column: 1 / -1;
                    margin-top: 10px;
                }
                
                /* Footer */
                .exam-footer {
                    margin-top: 50px;
                    padding-top: 20px;
                    border-top: 1px solid #bdc3c7;
                    text-align: center;
                    font-size: 12px;
                    color: #7f8c8d;
                    display: flex;
                    justify-content: space-between;
                    align-items: center;
                }
                
                .footer-left, .footer-right {
                    font-size: 11px;
                    color: #95a5a6;
                }
                
                /* Print Styles */
                @media print {
                    body {
                        padding: 20px;
                        font-size: 12pt;
                    }
                    
                    .exam-container {
                        max-width: none;
                    }
                    
                    .question {
                        page-break-inside: avoid;
                        break-inside: avoid;
                    }
                    
                    .university-header {
                        page-break-after: avoid;
                    }
                    
                    .instructions-section {
                        page-break-after: avoid;
                    }
                }
                
                /* Instructor Answer Indicators */
                .correct-answer {
                    background-color: #d4edda !important;
                    border-color: #28a745 !important;
                }
                
                .correct-answer .option-checkbox {
                    background-color: #28a745;
                    border-color: #28a745;
                    position: relative;
                }
                
                .correct-answer .option-checkbox::after {
                    content: '✓';
                    color: white;
                    font-size: 10px;
                    position: absolute;
                    top: -2px;
                    left: 1px;
                }
                
                .answer-key {
                    margin-top: 10px;
                    padding: 8px 12px;
                    background-color: #e8f4fd;
                    border: 1px solid #3498db;
                    border-radius: 4px;
                    font-size: 13px;
                    color: #2980b9;
                }
                
                .answer-key strong {
                    color: #1565C0;
                }
                
                
                .option {
                    margin: 12px 0;
                    padding: 12px 15px;
                    background: #f9fafb;
                    border: 2px solid #e5e7eb;
                    border-radius: 10px;
                    display: flex;
                    align-items: center;
                    cursor: pointer;
                    transition: all 200ms ease;
                }
            </style>
        </head>
        <body>
            <div class="exam-container">
                <!-- University Header -->
                <div class="university-header">
                    <div class="logo-placeholder">
                        <!-- Logo placeholder will be replaced with actual logo if available -->
                        {{ logo_image_placeholder }}
                    </div>
                    <div class="header-info">
                        <div class="university-name">{{ branding.university_name|default:"UNIVERSITY NAME" }}</div>
                        <div class="faculty-name">{{ branding.faculty|default:"FACULTY NAME" }}</div>
                        <div class="department-name">{{ branding.department|default:"DEPARTMENT NAME" }}</div>
                        <div class="course-info">{{ branding.course|default:"COURSE NAME" }}</div>
                        <div class="academic-info">{{ branding.academic_year|default:"ACADEMIC YEAR" }} "“ {{ branding.semester|default:"SEMESTER" }}</div>
                    </div>
                </div>
                
                <!-- Exam Title -->
                <div class="exam-title">
                    <div class="exam-type">{{ quiz_data.content_type|default:"EXAM"|upper }}</div>
                </div>
                
                <!-- Exam Metadata -->
                <div class="exam-metadata">
                    <div class="metadata-item">
                        <div class="metadata-label">INSTRUCTOR:</div>
                        <div class="metadata-value">{{ branding.instructor|default:"Instructor Name" }}</div>
                    </div>
                    <div class="metadata-item">
                        <div class="metadata-label">DATE:</div>
                        <div class="metadata-value">{{ branding.exam_date|default:"Exam Date" }}</div>
                    </div>
                    <div class="metadata-item">
                        <div class="metadata-label">DURATION:</div>
                        <div class="metadata-value">{{ quiz_data.estimated_duration|default:"2 hours" }}</div>
                    </div>
                    <div class="metadata-item">
                        <div class="metadata-label">TOTAL:</div>
                        <div class="metadata-value">{{ quiz_data.total_points|default:quiz_data.questions|length }} points</div>
                    </div>
                </div>
                
                <!-- Student Information Section -->
                <div class="student-info-section">
                    <h3 class="student-info-title">STUDENT INFORMATION</h3>
                    <div class="student-fields">
                        <!-- These fields will be conditionally rendered based on branding settings -->
                        {{ student_info_fields_placeholder }}
                    </div>
                </div>
                
                <!-- Instructions -->
                <div class="instructions-section">
                    <div class="instructions-title">Instructions:</div>
                    <ul class="instructions-list">
                        <li>Read each question carefully</li>
                        <li>Answer clearly in the space provided</li>
                        <li>Calculators are permitted unless otherwise noted</li>
                        <li>Show all work for calculation problems</li>
                    </ul>
                </div>
                
                <!-- Questions -->
                {{ questions_placeholder }}
                
                <!-- Footer -->
                <div class="exam-footer">
                    <div class="footer-left">
                        {{ branding.academic_year|default:"Academic Year" }} | {{ branding.semester|default:"Semester" }}
                    </div>
                    <div class="footer-center">
                        {{ branding.university_name|default:"University" }} "“ Official {{ quiz_data.content_type|default:"Exam" }}
                    </div>
                    <div class="footer-right">
                        {{ branding.department|default:"Department" }}
                    </div>
                </div>
            </div>
        </body>
        </html>
        '''
        
        # Add instructor-specific CSS if showing answers
        if show_answers:
            instructor_css = '''
            <style>
                /* Instructor version - correct answer styling */
                .correct-option {
                    background: #ecfdf5 !important;
                    border: 2px solid #10b981 !important;
                    font-weight: 600;
                }
                
                .correct-option .option-letter {
                    background: #10b981;
                    color: white;
                    padding: 2px 8px;
                    border-radius: 4px;
                }
                
                /* Expected answer styling for instructor version */
                .expected-answer {
                    background: #f0f9ff;
                    border: 1px solid #0ea5e9;
                    border-radius: 6px;
                    padding: 8px 12px;
                    margin-top: 10px;
                    color: #0c4a6e;
                    font-size: 14px;
                }
                
                .expected-answer strong {
                    color: #0369a1;
                }
            </style>
            '''
            # Insert the instructor CSS before </head>
            html_template = html_template.replace('</head>', instructor_css + '\n        </head>')
        
        # Simple template rendering (replace all placeholders)
        html = html_template
        
        # Basic replacements
        html = html.replace('{{ quiz_data.title }}', quiz_data.get('title', 'Quiz'))
        html = html.replace('{{ quiz_data.description }}', quiz_data.get('description', ''))
        html = html.replace('{{ export_date }}', datetime.now().strftime('%B %d, %Y'))
        
        # Comprehensive branding replacements
        html = html.replace('{{ branding.university_name|default:"UNIVERSITY NAME" }}', branding.get('university_name', 'UNIVERSITY NAME') if branding else 'UNIVERSITY NAME')
        html = html.replace('{{ branding.faculty|default:"FACULTY NAME" }}', branding.get('faculty', 'FACULTY NAME') if branding else 'FACULTY NAME')
        html = html.replace('{{ branding.department|default:"DEPARTMENT NAME" }}', branding.get('department', 'DEPARTMENT NAME') if branding else 'DEPARTMENT NAME')
        html = html.replace('{{ branding.course|default:"COURSE NAME" }}', branding.get('course', 'COURSE NAME') if branding else 'COURSE NAME')
        html = html.replace('{{ branding.academic_year|default:"ACADEMIC YEAR" }}', branding.get('academic_year', 'ACADEMIC YEAR') if branding else 'ACADEMIC YEAR')
        html = html.replace('{{ branding.semester|default:"SEMESTER" }}', branding.get('semester', 'SEMESTER') if branding else 'SEMESTER')
        html = html.replace('{{ branding.instructor|default:"Instructor Name" }}', branding.get('instructor', 'Instructor Name') if branding else 'Instructor Name')
        html = html.replace('{{ branding.exam_date|default:"Exam Date" }}', branding.get('exam_date', 'Exam Date') if branding else 'Exam Date')
        
        # More comprehensive template cleanup
        import re
        
        # Replace branding placeholders
        if branding and branding.get('institution_name'):
            html = re.sub(r'{% if branding\.institution_name %}\s*<div class="institution">{{ branding\.institution_name }}</div>\s*{% endif %}', 
                         f'<div class="institution">{branding["institution_name"]}</div>', html, flags=re.DOTALL)
        else:
            html = re.sub(r'{% if branding\.institution_name %}.*?{% endif %}', '', html, flags=re.DOTALL)
            
        if branding and branding.get('department'):
            html = re.sub(r'{% if branding\.department %}\s*<div class="department">{{ branding\.department }}</div>\s*{% endif %}', 
                         f'<div class="department">{branding["department"]}</div>', html, flags=re.DOTALL)
        else:
            html = re.sub(r'{% if branding\.department %}.*?{% endif %}', '', html, flags=re.DOTALL)
        
        # Replace description placeholder more comprehensively
        description_placeholder = '''{% if quiz_data.description %}
                    <p class="description">{{ quiz_data.description }}</p>
                    {% endif %}'''
        if quiz_data.get('description'):
            html = html.replace(description_placeholder, f'<p class="description">{quiz_data["description"]}</p>')
        else:
            html = html.replace(description_placeholder, '')
        
        # Replace quiz data placeholders
        html = html.replace('{{ quiz_data.estimated_duration|default:"Not specified" }}', 
                           quiz_data.get('estimated_duration', 'Not specified'))
        html = html.replace('{{ quiz_data.total_points|default:quiz_data.questions|length }}', 
                           str(quiz_data.get('total_points', len(quiz_data.get('questions', [])))))
        
        # Add questions - clean professional format without type labels
        questions_html = ""
        for i, question in enumerate(quiz_data.get('questions', []), 1):
            question_type = question.get('type', 'multiple_choice')
            
            # Escape HTML special characters in question content
            import html as html_escape_module
            question_text = html_escape_module.escape(question.get('question', '')) if question.get('question') else ''
            
            # University-style question formatting
            q_html = f"""
                <div class="question">
                    <div class="question-header">
                        <div class="question-number">{i}.</div>
                        <div class="question-points">({question.get('points', 1)} point{'s' if question.get('points', 1) != 1 else ''})</div>
                    </div>
                    <div class="question-text">{question_text}</div>
            """
            
            # Handle different question types
            if question_type == 'multiple_choice':
                if question.get('options') and len(question['options']) > 0:
                    # Remove duplicate options and support up to 5 options (A-E)
                    unique_options = list(dict.fromkeys(question['options']))[:5]
                    correct_answer = question.get('correct_answer', '').upper() if show_answers else ''
                    q_html += '<div class="mc-options">'
                    for j, option in enumerate(unique_options):
                        option_letter = chr(65 + j)  # A, B, C, D, E
                        escaped_option = html_escape_module.escape(str(option)) if option else ''
                        
                        # Only show correct answer indicators if show_answers is True
                        correct_class = ' correct-answer' if show_answers and correct_answer == option_letter else ''
                        
                        q_html += f'''
                            <div class="mc-option{correct_class}">
                                <div class="option-checkbox"></div>
                                <div class="option-letter">{option_letter}.</div>
                                <div class="option-text">{escaped_option}</div>
                            </div>
                        '''
                    q_html += '</div>'
                    
                    # Add explanation for instructor version
                    if show_answers and question.get('explanation'):
                        explanation = html_escape_module.escape(str(question.get('explanation', '')))
                        q_html += f'<div class="answer-key"><strong>Explanation:</strong> {explanation}</div>'
                else:
                    # Fallback if no options provided - support up to E
                    q_html += '<div class="mc-options">'
                    for j, letter in enumerate(['A', 'B', 'C', 'D', 'E']):
                        q_html += f'''
                            <div class="mc-option">
                                <div class="option-checkbox"></div>
                                <div class="option-letter">{letter}.</div>
                                <div class="option-text">___________</div>
                            </div>
                        '''
                    q_html += '</div>'
            
            elif question_type == 'true_false':
                true_label = 'True'
                false_label = 'False'
                correct_answer = str(question.get('correct_answer', '')).lower() if show_answers else ''
                
                true_class = ' correct-answer' if show_answers and correct_answer == 'true' else ''
                false_class = ' correct-answer' if show_answers and correct_answer == 'false' else ''
                
                q_html += f'''
                    <div class="tf-options">
                        <div class="tf-option{true_class}">
                            <div class="option-checkbox"></div>
                            <span>{true_label}</span>
                        </div>
                        <div class="tf-option{false_class}">
                            <div class="option-checkbox"></div>
                            <span>{false_label}</span>
                        </div>
                    </div>
                '''
            
            elif question_type == 'short_answer':
                q_html += '<div class="answer-space"></div>'
                
                # Show expected answer and explanation for instructor version
                if show_answers and question.get('correct_answer'):
                    expected_answer = html_escape_module.escape(str(question['correct_answer']))
                    explanation = html_escape_module.escape(str(question.get('explanation', '')))
                    q_html += f'<div class="answer-key"><strong>Expected Answer:</strong> {expected_answer}'
                    if explanation:
                        q_html += f'<br><strong>Explanation:</strong> {explanation}'
                    q_html += '</div>'
            
            elif question_type == 'fill_blank':
                # Add blank spaces for fill in the blank questions
                q_html += '<div class="fill-blank">'
                q_html += '<div class="blank-space"></div>'
                q_html += '</div>'
                
                # Show expected answer for instructor version
                if show_answers and question.get('correct_answer'):
                    expected_answer = html_escape_module.escape(str(question['correct_answer']))
                    q_html += f'<div class="answer-key"><strong>Expected Answer:</strong> {expected_answer}</div>'
            
            elif question_type == 'essay':
                # Use the essay-space styling for lined writing area
                q_html += '<div class="essay-space"></div>'
                
                # Show expected answer for instructor version
                if show_answers and question.get('correct_answer'):
                    expected_answer = html_escape_module.escape(str(question['correct_answer']))
                    q_html += f'<div class="answer-key"><strong>Expected Answer:</strong> {expected_answer}</div>'
            
            # Add catch-all for any unhandled question type
            else:
                answer_label = 'Cevap:' if 'tr' in str(quiz_data.get('language', '')).lower() else 'Answer:'
                answer_space = '<span class="answer-space"></span>'
                
                # Show expected answer for instructor version
                if show_answers and question.get('correct_answer'):
                    expected_answer = html_escape_module.escape(str(question['correct_answer']))
                    answer_space += f'<div class="expected-answer"><strong>Expected:</strong> {expected_answer}</div>'
                
                q_html += f'<div class="answer-container">{answer_label} {answer_space}</div>'
            
            q_html += '</div>'
            questions_html += q_html
        
        # Replace quiz data placeholders for content type and other fields
        html = html.replace('{{ quiz_data.content_type|default:"EXAM"|upper }}', quiz_data.get('content_type', 'QUIZ').upper())
        html = html.replace('{{ quiz_data.estimated_duration|default:"2 hours" }}', quiz_data.get('estimated_duration', '2 hours'))
        html = html.replace('{{ quiz_data.total_points|default:quiz_data.questions|length }} points', f"{quiz_data.get('total_points', len(quiz_data.get('questions', [])))} points")
        
        # Footer replacements
        html = html.replace('{{ branding.university_name|default:"University" }}', branding.get('university_name', 'University') if branding else 'University')
        html = html.replace('{{ quiz_data.content_type|default:"Exam" }}', quiz_data.get('content_type', 'Quiz'))
        
        # Generate student information fields based on branding settings
        student_info_html = ""
        if branding and 'student_info' in branding:
            student_info = branding['student_info']
            student_fields = []
            
            if student_info.get('include_student_name', True):
                student_fields.append('<div class="student-field"><div class="student-field-label">FULL NAME:</div><div class="student-field-line"></div></div>')
            
            if student_info.get('include_student_id', True):
                student_fields.append('<div class="student-field"><div class="student-field-label">STUDENT ID:</div><div class="student-field-line"></div></div>')
            
            if student_info.get('include_date_field', False):
                student_fields.append('<div class="student-field"><div class="student-field-label">DATE:</div><div class="student-field-line"></div></div>')
            
            if student_info.get('include_signature', True):
                student_fields.append('<div class="student-field signature-field"><div class="student-field-label">SIGNATURE:</div><div class="student-field-line"></div></div>')
            
            if student_fields:
                student_info_html = '\n'.join(student_fields)
        else:
            # Default student fields if no specific configuration
            student_info_html = '''
                <div class="student-field"><div class="student-field-label">FULL NAME:</div><div class="student-field-line"></div></div>
                <div class="student-field"><div class="student-field-label">STUDENT ID:</div><div class="student-field-line"></div></div>
                <div class="student-field signature-field"><div class="student-field-label">SIGNATURE:</div><div class="student-field-line"></div></div>
            '''
        
        # Replace logo placeholder with actual logo if available
        # Prefer embedded base64 so exported HTML works offline and in ZIP downloads.
        logo_html = 'UNIVERSITY<br/>LOGO<br/><small>(Upload in export form)</small>'  # Default placeholder
        if branding:
            try:
                logo_resolved = _read_logo_bytes_from_branding(branding)
                if logo_resolved:
                    import base64
                    logo_bytes, mime = logo_resolved
                    logo_bytes, mime = _prepare_logo_for_embedding(logo_bytes, mime, max_px=256)
                    b64 = base64.b64encode(logo_bytes).decode('ascii')
                    logo_html = f'<img src="data:{mime};base64,{b64}" alt="University Logo" style="max-width: 80px; max-height: 80px; object-fit: contain;"/>'
                    logger.info("Added logo to HTML export using embedded data URL")
                elif branding.get('logo_url'):
                    # Fallback to URL when we cannot resolve bytes
                    logo_url = branding.get('logo_url')
                    if isinstance(logo_url, str):
                        if not logo_url.startswith('http') and not logo_url.startswith('/'):
                            logo_url = '/media/' + logo_url.lstrip('/')
                        logo_html = f'<img src="{logo_url}" alt="University Logo" style="max-width: 80px; max-height: 80px; object-fit: contain;"/>'
                        logger.info(f"Added logo to HTML export using URL: {logo_url}")
            except Exception as e:
                logger.warning(f"Could not add logo to HTML export: {e}")
        
        html = html.replace('{{ logo_image_placeholder }}', logo_html)
        
        # Replace student info placeholder
        html = html.replace('{{ student_info_fields_placeholder }}', student_info_html)
        
        # Replace questions placeholder with generated content
        html = html.replace('{{ questions_placeholder }}', questions_html)
        
        # Add watermark if specified
        if branding and branding.get('watermark'):
            watermark_text = branding['watermark']
            if watermark_text.strip():
                # Add watermark CSS and HTML
                watermark_css = f'''
                <style>
                    .watermark {{
                        position: fixed;
                        top: 50%;
                        left: 50%;
                        transform: translate(-50%, -50%) rotate(45deg);
                        font-size: 80px;
                        color: rgba(200, 200, 200, 0.2);
                        font-weight: bold;
                        font-family: Arial, sans-serif;
                        z-index: -1;
                        pointer-events: none;
                        user-select: none;
                    }}
                    @media print {{
                        .watermark {{
                            position: fixed !important;
                        }}
                    }}
                </style>
                '''
                
                watermark_html = f'<div class="watermark">{watermark_text}</div>'
                
                # Insert watermark CSS before </head>
                html = html.replace('</head>', watermark_css + '\n        </head>')
                
                # Insert watermark HTML after <body>
                html = html.replace('<body>', '<body>\n            ' + watermark_html)
        
        return html


class ZIPExporter:
    """Service for creating ZIP archives with multiple formats"""
    
    def __init__(self):
        self.pdf_exporter = PDFExporter()
        self.html_exporter = HTMLExporter()
        if DOCX_AVAILABLE:
            self.docx_exporter = DOCXExporter()
    
    def export_complete_package(self, content_data: Dict[str, Any], 
                              versions: List[str] = None,
                              formats: List[str] = None,
                              branding: Dict[str, Any] = None) -> io.BytesIO:
        """
        Create a complete export package with multiple formats and versions
        
        Args:
            content_data: The quiz/exam data
            versions: List of version letters (e.g., ['A', 'B', 'C'])
            formats: List of formats to include ('pdf', 'docx', 'html')
            branding: Branding information
            
        Returns:
            BytesIO buffer containing the ZIP file
        """
        buffer = io.BytesIO()
        versions = versions or ['A']
        formats = formats or ['pdf']
        
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            content_type = content_data.get('type', 'quiz')
            base_title = content_data.get('title', 'Content')
            
            for version in versions:
                # Create version-specific data
                version_data = self._create_version_data(content_data, version)
                
                for format_type in formats:
                    try:
                        if format_type == 'pdf':
                            # Export questions
                            if content_type == 'quiz':
                                pdf_buffer = self.pdf_exporter.export_quiz(version_data, branding)
                            else:
                                pdf_buffer = self.pdf_exporter.export_exam(version_data, branding)
                            
                            zipf.writestr(f'{base_title}_Version_{version}.pdf', pdf_buffer.getvalue())
                            
                            # Export answer key
                            answer_key_buffer = self.pdf_exporter.export_answer_key(version_data, branding)
                            zipf.writestr(f'{base_title}_Version_{version}_Answer_Key.pdf', answer_key_buffer.getvalue())
                        
                        elif format_type == 'docx' and DOCX_AVAILABLE:
                            if content_type == 'quiz':
                                docx_buffer = self.docx_exporter.export_quiz(version_data, branding)
                            else:
                                docx_buffer = self.docx_exporter.export_exam(version_data, branding)
                            
                            zipf.writestr(f'{base_title}_Version_{version}.docx', docx_buffer.getvalue())
                        
                        elif format_type == 'html':
                            html_content = self.html_exporter.export_quiz(version_data, branding)
                            zipf.writestr(f'{base_title}_Version_{version}.html', html_content.encode('utf-8'))
                    
                    except Exception as e:
                        logger.error(f"Error exporting {format_type} for version {version}: {str(e)}")
                        continue
            
            # Add metadata file
            metadata = {
                'title': base_title,
                'content_type': content_type,
                'versions': versions,
                'formats': formats,
                'export_date': datetime.now().isoformat(),
                'branding': branding or {}
            }
            zipf.writestr('metadata.json', json.dumps(metadata, indent=2))
        
        buffer.seek(0)
        return buffer
    
    def _create_version_data(self, content_data: Dict[str, Any], version: str) -> Dict[str, Any]:
        """Create version-specific content data"""
        import random
        
        version_data = content_data.copy()
        version_data['title'] = f"{content_data.get('title', 'Content')} - Version {version}"
        
        # Shuffle questions for different versions
        if 'questions' in version_data:
            questions = version_data['questions'].copy()
            random.seed(ord(version))  # Use version letter as seed for reproducibility
            random.shuffle(questions)
            version_data['questions'] = questions
        elif 'sections' in version_data:
            # Handle exam format
            for section in version_data['sections']:
                if 'questions' in section:
                    questions = section['questions'].copy()
                    random.seed(ord(version))
                    random.shuffle(questions)
                    section['questions'] = questions
        
        return version_data


class ExportService:
    """Main export service that coordinates all export types"""
    
    def __init__(self):
        self.pdf_exporter = PDFExporter()
        self.html_exporter = HTMLExporter()
        self.zip_exporter = ZIPExporter()
        
        if DOCX_AVAILABLE:
            self.docx_exporter = DOCXExporter()
    
    def export_content(self, content_data: Dict[str, Any], 
                      export_format: str,
                      branding: Dict[str, Any] = None,
                      include_answer_key: bool = True,
                      versions: List[str] = None) -> Dict[str, Any]:
        """
        Main export method that handles all formats
        
        Args:
            content_data: The content to export
            export_format: 'pdf', 'docx', 'html', or 'zip'
            branding: Branding information
            include_answer_key: Whether to include answer key
            versions: Version letters for multi-version exports
            
        Returns:
            Dict with export results
        """
        try:
            if export_format == 'pdf':
                return self._export_pdf(content_data, branding, include_answer_key)
            elif export_format == 'docx':
                return self._export_docx(content_data, branding)
            elif export_format == 'html':
                return self._export_html(content_data, branding)
            elif export_format == 'json':
                return self._export_json(content_data, branding)
            elif export_format == 'zip':
                return self._export_zip(content_data, branding, versions or ['A', 'B', 'C'])
            else:
                return {
                    'success': False,
                    'error': f'Unsupported export format: {export_format}'
                }
        
        except Exception as e:
            logger.error(f"Export error: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _export_pdf(self, content_data: Dict[str, Any], branding: Dict[str, Any], include_answer_key: bool) -> Dict[str, Any]:
        """Export to PDF format"""
        content_type = content_data.get('type', 'quiz')
        
        if content_type == 'quiz':
            buffer = self.pdf_exporter.export_quiz(content_data, branding)
        else:
            buffer = self.pdf_exporter.export_exam(content_data, branding)
        
        result = {
            'success': True,
            'file_data': buffer.getvalue(),
            'filename': f"{content_data.get('title', 'content')}.pdf",
            'content_type': 'application/pdf'
        }
        
        if include_answer_key:
            answer_key_buffer = self.pdf_exporter.export_answer_key(content_data, branding)
            result['answer_key_data'] = answer_key_buffer.getvalue()
            result['answer_key_filename'] = f"{content_data.get('title', 'content')}_answer_key.pdf"
        
        return result
    
    def _export_docx(self, content_data: Dict[str, Any], branding: Dict[str, Any]) -> Dict[str, Any]:
        """Export to DOCX format"""
        if not DOCX_AVAILABLE:
            return {
                'success': False,
                'error': 'DOCX export not available. Install python-docx.'
            }
        
        content_type = content_data.get('type', 'quiz')
        
        if content_type == 'quiz':
            buffer = self.docx_exporter.export_quiz(content_data, branding)
        else:
            buffer = self.docx_exporter.export_exam(content_data, branding)
        
        return {
            'success': True,
            'file_data': buffer.getvalue(),
            'filename': f"{content_data.get('title', 'content')}.docx",
            'content_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
    
    def _export_html(self, content_data: Dict[str, Any], branding: Dict[str, Any]) -> Dict[str, Any]:
        """Export to HTML format"""
        html_content = self.html_exporter.export_quiz(content_data, branding)
        
        return {
            'success': True,
            'file_data': html_content.encode('utf-8'),
            'filename': f"{content_data.get('title', 'content')}.html",
            'content_type': 'text/html'
        }
    
    def _export_json(self, content_data: Dict[str, Any], branding: Dict[str, Any]) -> Dict[str, Any]:
        """Export to JSON format"""
        # Add metadata to content
        export_data = {
            'export_metadata': {
                'export_date': datetime.now().isoformat(),
                'format': 'json',
                'branding': branding
            },
            'content': content_data
        }
        
        json_content = json.dumps(export_data, indent=2, default=str)
        
        return {
            'success': True,
            'file_data': json_content.encode('utf-8'),
            'filename': f"{content_data.get('title', 'content')}.json",
            'content_type': 'application/json'
        }
    
    def _export_zip(self, content_data: Dict[str, Any], branding: Dict[str, Any], versions: List[str]) -> Dict[str, Any]:
        """Export to ZIP format with multiple versions"""
        buffer = self.zip_exporter.export_complete_package(
            content_data, 
            versions=versions,
            formats=['pdf', 'html'] + (['docx'] if DOCX_AVAILABLE else []),
            branding=branding
        )
        
        return {
            'success': True,
            'file_data': buffer.getvalue(),
            'filename': f"{content_data.get('title', 'content')}_complete_package.zip",
            'content_type': 'application/zip'
        }
    
    def export_generation(self, export_job) -> Dict[str, Any]:
        """
        Export an AI generation to the specified format
        
        Args:
            export_job: ExportJob instance containing export configuration
            
        Returns:
            Dict with export results
        """
        try:
            # Get generation data
            generation = export_job.generation
            
            # Prepare content data
            content_data = self._prepare_generation_data(generation)
            
            # Prepare branding data
            branding = export_job.branding_settings or {}
            if export_job.watermark:
                branding['watermark'] = export_job.watermark
            
            # Ensure logo information is included if university_logo exists
            if export_job.university_logo:
                try:
                    import os
                    from django.conf import settings
                    logo_path = export_job.university_logo.path
                    logo_url = export_job.university_logo.url
                    
                    if os.path.exists(logo_path):
                        branding['logo_path'] = logo_path
                        branding['logo_url'] = logo_url
                        branding['logo_filename'] = export_job.university_logo.name
                        branding['has_logo'] = True
                        logger.info(f"Logo included in branding: {logo_path}")
                    else:
                        # Try alternative path resolution
                        if logo_url.startswith('/media/'):
                            alt_path = os.path.join(settings.MEDIA_ROOT, logo_url.replace('/media/', ''))
                            if os.path.exists(alt_path):
                                branding['logo_path'] = alt_path
                                branding['logo_url'] = logo_url
                                branding['has_logo'] = True
                                logger.info(f"Logo found at alternative path: {alt_path}")
                except Exception as e:
                    logger.warning(f"Could not add logo to branding in export_generation: {e}")
            
            # Export using the main export method
            result = self.export_content(
                content_data=content_data,
                export_format=export_job.export_format,
                branding=branding,
                include_answer_key=export_job.include_answer_key
            )
            
            if result['success']:
                # Save the exported file
                file_content = ContentFile(result['file_data'])
                export_job.generated_file.save(
                    result['filename'],
                    file_content,
                    save=False  # Don't save yet, we need to set file_size first
                )
                export_job.file_size = len(result['file_data'])
                export_job.save()  # Save the export_job with file_size
                export_job.mark_completed()
                
                # Save answer key if available
                if 'answer_key_data' in result:
                    # Create a separate export for answer key or store as metadata
                    export_job.parameters['answer_key_available'] = True
                    export_job.save()
            else:
                export_job.mark_error(result.get('error', 'Unknown export error'))
            
            return result
            
        except Exception as e:
            error_msg = f"Export generation failed: {str(e)}"
            logger.error(error_msg)
            export_job.mark_error(error_msg)
            return {
                'success': False,
                'error': error_msg
            }
    
    def export_with_versions(self, export_job, num_versions: int = 3) -> Dict[str, Any]:
        """
        Export an AI generation with multiple versions (A, B, C)
        
        Args:
            export_job: ExportJob instance containing export configuration
            num_versions: Number of versions to create
            
        Returns:
            Dict with export results
        """
        try:
            # Get generation data
            generation = export_job.generation
            
            # Prepare content data
            content_data = self._prepare_generation_data(generation)
            
            # Prepare branding data
            branding = export_job.branding_settings or {}
            if export_job.watermark:
                branding['watermark'] = export_job.watermark
            
            # Create version letters
            version_letters = [chr(65 + i) for i in range(num_versions)]  # A, B, C, etc.
            
            if export_job.export_format == 'zip':
                # Export as ZIP with multiple versions
                result = self._export_zip(content_data, branding, version_letters)
            else:
                # For non-ZIP formats, create individual version files
                result = self._export_individual_versions(
                    export_job, content_data, branding, version_letters
                )
            
            if result['success']:
                # Save the main export file
                file_content = ContentFile(result['file_data'])
                export_job.generated_file.save(
                    result['filename'],
                    file_content,
                    save=False
                )
                export_job.file_size = len(result['file_data'])
                export_job.save()  # Save the export_job with file_size
                export_job.mark_completed()
            else:
                export_job.mark_error(result.get('error', 'Unknown export error'))
            
            return result
            
        except Exception as e:
            error_msg = f"Export with versions failed: {str(e)}"
            logger.error(error_msg)
            export_job.mark_error(error_msg)
            return {
                'success': False,
                'error': error_msg
            }
    
    def _prepare_generation_data(self, generation) -> Dict[str, Any]:
        """
        Convert AI generation to export-ready data format
        
        Args:
            generation: AIGeneration instance
            
        Returns:
            Dict containing structured content data
        """
        content_data = {
            'title': generation.title,
            'description': generation.description or '',
            'type': generation.content_type,
            'created_at': generation.created_at,
        }
        
        # Handle generated content
        if generation.generated_content:
            generated = generation.generated_content
            
            # Extract questions if available
            if 'questions' in generated:
                content_data['questions'] = generated['questions']
            elif 'sections' in generated:
                content_data['sections'] = generated['sections']
            
            # Extract metadata
            content_data.update({
                'total_points': generated.get('total_points', 0),
                'estimated_duration': generated.get('estimated_duration', ''),
            })
        
        # Handle individual questions from database
        if hasattr(generation, 'questions') and generation.questions.exists():
            questions = []
            for q in generation.questions.all():
                question_data = {
                    'id': q.id,
                    'type': q.question_type,
                    'question': q.question_text,
                    'points': q.points,
                    'difficulty': q.difficulty,
                    'correct_answer': q.correct_answer,
                    'explanation': q.explanation or '',
                }
                
                # Add options for multiple choice questions
                if q.question_type == 'multiple_choice' and q.options:
                    question_data['options'] = q.options
                
                questions.append(question_data)
            
            content_data['questions'] = questions
            content_data['total_points'] = sum(q.points for q in generation.questions.all())
        
        return content_data
    
    def _export_individual_versions(self, export_job, content_data: Dict[str, Any], 
                                   branding: Dict[str, Any], version_letters: List[str]) -> Dict[str, Any]:
        """
        Export individual version files and create ExportVersion records
        
        Args:
            export_job: ExportJob instance
            content_data: Content to export
            branding: Branding information
            version_letters: List of version letters
            
        Returns:
            Dict with export results (returns the first version as main file)
        """
        from .models import ExportVersion
        import random
        
        results = []
        main_result = None
        
        for i, version_letter in enumerate(version_letters):
            # Create version-specific data
            version_data = content_data.copy()
            version_data['title'] = f"{content_data.get('title', 'Content')} - Version {version_letter}"
            
            # Randomize question order for versions
            if 'questions' in version_data:
                questions = version_data['questions'].copy()
                random.seed(ord(version_letter))  # Reproducible randomization
                random.shuffle(questions)
                version_data['questions'] = questions
            
            # Export this version
            result = self.export_content(
                content_data=version_data,
                export_format=export_job.export_format,
                branding=branding,
                include_answer_key=export_job.include_answer_key
            )
            
            if result['success']:
                # Create ExportVersion record
                export_version = ExportVersion.objects.create(
                    export_job=export_job,
                    version_letter=version_letter,
                    file_size=len(result['file_data']),
                    variations={'randomized_order': True}
                )
                
                # Save version file
                version_filename = f"{content_data.get('title', 'content')}_Version_{version_letter}.{export_job.export_format}"
                file_content = ContentFile(result['file_data'])
                export_version.generated_file.save(
                    version_filename,
                    file_content,
                    save=True
                )
                
                # Use first version as main result
                if i == 0:
                    main_result = result
                    main_result['filename'] = version_filename
                
                results.append(result)
        
        return main_result or {
            'success': False,
            'error': 'Failed to create any versions'
        }

